import streamlit as st
from openai import OpenAI
import weaviate
from bs4 import BeautifulSoup
import requests
import PyPDF2
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import io
from docx import Document as DocxDocument
from fpdf import FPDF
from typing import List, Dict, Tuple, Any, Optional
import os
from datetime import datetime
import pandas as pd
import time
import threading
from queue import Queue
import hashlib
import pickle
import logging
import concurrent.futures
import sys
import psutil
from functools import lru_cache
import validators
from collections import deque

# LangChain imports
from langchain.llms import OpenAI as LangChainOpenAI
from langchain.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Weaviate as LangChainWeaviate
from langchain.chains import (
    ConversationalRetrievalChain,
    LLMChain,
    SequentialChain,
    RetrievalQA
)
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate
from langchain.document_loaders import (
    PyPDFLoader,
    TextLoader,
    Docx2txtLoader,
    UnstructuredURLLoader
)
from langchain.callbacks import StreamlitCallbackHandler
from langchain.agents import (
    Tool,
    AgentExecutor,
    ConversationalChatAgent,
    ZeroShotAgent
)

# Initialize OpenAI and Weaviate clients
client_openai = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
client = weaviate.Client(
    url=st.secrets["WEAVIATE_URL"],
    auth_client_secret=weaviate.AuthApiKey(api_key=st.secrets["WEAVIATE_API_KEY"]),
    additional_headers={"X-OpenAI-Api-Key": st.secrets["OPENAI_API_KEY"]}
)

# Initialize LangChain components
llm = LangChainOpenAI(
    temperature=0.3,
    model_name="gpt-3.5-turbo",
    streaming=True,
    openai_api_key=st.secrets["OPENAI_API_KEY"]
)

embeddings = OpenAIEmbeddings(
    openai_api_key=st.secrets["OPENAI_API_KEY"]
)

# Define processing queue
if "processing_queue" not in st.session_state:
    st.session_state.processing_queue = Queue()

# Initialize custom logger
class CustomLogger:
    def __init__(self):
        self.logger = logging.getLogger('COGNITEXT')
        self.logger.setLevel(logging.INFO)
        
        # File handler
        fh = logging.FileHandler('cognitext.log')
        fh.setLevel(logging.INFO)
        
        # Console handler
        ch = logging.StreamHandler()
        ch.setLevel(logging.ERROR)
        
        # Formatter
        formatter = logging.Formatter(
            '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        fh.setFormatter(formatter)
        ch.setFormatter(formatter)
        
        self.logger.addHandler(fh)
        self.logger.addHandler(ch)
        
    def log_error(self, error: str, context: Dict = None):
        self.logger.error(f"Error: {error}, Context: {context}")
        
    def log_info(self, message: str):
        self.logger.info(message)
        
    def log_warning(self, message: str):
        self.logger.warning(message)

logger = CustomLogger()

# Custom exceptions
class CognitextException(Exception):
    """Base exception class for COGNITEXT."""
    pass

class DocumentProcessingError(CognitextException):
    """Raised when document processing fails."""
    pass

class DatabaseError(CognitextException):
    """Raised when database operations fail."""
    pass

class APIError(CognitextException):
    """Raised when API calls fail."""
    pass

# Caching system
class CacheManager:
    def __init__(self):
        self.cache_dir = ".cache"
        os.makedirs(self.cache_dir, exist_ok=True)
        
    def get_cache_key(self, data):
        """Generate a unique cache key."""
        return hashlib.md5(str(data).encode()).hexdigest()
        
    def cache_exists(self, key):
        return os.path.exists(os.path.join(self.cache_dir, key))
        
    def get_cache(self, key):
        try:
            with open(os.path.join(self.cache_dir, key), 'rb') as f:
                return pickle.load(f)
        except:
            return None
            
    def set_cache(self, key, data):
        try:
            with open(os.path.join(self.cache_dir, key), 'wb') as f:
                pickle.dump(data, f)
        except:
            pass

cache_manager = CacheManager()

# Memory management
class MemoryManager:
    def __init__(self, max_memory_mb: int = 1000):
        self.max_memory = max_memory_mb * 1024 * 1024
        
    def check_memory(self):
        """Check current memory usage."""
        process = psutil.Process(os.getpid())
        return process.memory_info().rss
        
    def can_process(self, estimated_size: int) -> bool:
        """Check if we can process given estimated size."""
        return self.check_memory() + estimated_size < self.max_memory

memory_manager = MemoryManager()

# Input validation
class InputValidator:
    @staticmethod
    def validate_url(url: str) -> bool:
        return bool(validators.url(url))
        
    @staticmethod
    def validate_file_type(file_type: str, allowed_types: List[str]) -> bool:
        return file_type.lower() in allowed_types
        
    @staticmethod
    def sanitize_text(text: str) -> str:
        """Remove potentially harmful characters."""
        import re
        return re.sub(r'[<>]', '', text)
        
    @staticmethod
    def validate_file_size(file_size: int, max_size_mb: int = 10) -> bool:
        return file_size <= max_size_mb * 1024 * 1024

input_validator = InputValidator()

# Rate limiting
class RateLimiter:
    def __init__(self, max_requests: int, time_window: int):
        self.max_requests = max_requests
        self.time_window = time_window
        self.requests = deque()
        
    def can_proceed(self) -> bool:
        now = time.time()
        
        # Remove old requests
        while self.requests and self.requests[0] <= now - self.time_window:
            self.requests.popleft()
            
        if len(self.requests) < self.max_requests:
            self.requests.append(now)
            return True
        return False
        
    def wait_time(self) -> float:
        if len(self.requests) < self.max_requests:
            return 0
        return self.requests[0] + self.time_window - time.time()

# Initialize rate limiters
api_limiter = RateLimiter(max_requests=60, time_window=60)  # 60 requests per minute
# Document Processing System
class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        
    def load_document(self, file_info: Dict) -> List[Dict]:
        """Load document using appropriate LangChain loader."""
        file_type = file_info['type']
        file_path = self._save_temp_file(file_info['bytes'], file_info['name'])
        
        try:
            if file_type == 'pdf':
                if file_info.get('use_ocr', False):
                    return self._process_with_ocr(file_info['bytes'])
                loader = PyPDFLoader(file_path)
            elif file_type == 'docx':
                loader = Docx2txtLoader(file_path)
            elif file_type == 'txt':
                loader = TextLoader(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
            documents = loader.load()
            splits = self.text_splitter.split_documents(documents)
            return splits
            
        finally:
            if os.path.exists(file_path):
                os.remove(file_path)
    
    def _process_with_ocr(self, pdf_bytes: bytes) -> List[Dict]:
        """Process PDF with OCR using pytesseract."""
        try:
            images = convert_pdf_to_images(pdf_bytes)
            texts = []
            for image in images:
                text = pytesseract.image_to_string(image)
                if text.strip():
                    texts.append(text)
            
            combined_text = "\n\n".join(texts)
            return self.text_splitter.split_text(combined_text)
        except Exception as e:
            logger.log_error(f"OCR processing error: {str(e)}")
            raise DocumentProcessingError("OCR processing failed")
    
    def load_url(self, url: str) -> List[Dict]:
        """Load content from URL using LangChain."""
        if not input_validator.validate_url(url):
            raise ValueError("Invalid URL provided")
            
        loader = UnstructuredURLLoader([url])
        documents = loader.load()
        return self.text_splitter.split_documents(documents)
        
    def _save_temp_file(self, content: bytes, filename: str) -> str:
        """Save content to temporary file."""
        temp_dir = "temp"
        os.makedirs(temp_dir, exist_ok=True)
        file_path = os.path.join(temp_dir, filename)
        
        with open(file_path, 'wb') as f:
            f.write(content)
        return file_path

# Vector Store Management
class VectorStoreManager:
    def __init__(self):
        self.vectorstore = LangChainWeaviate(
            client=client,
            index_name="Document",
            text_key="content",
            embedding=embeddings,
            by_text=False
        )
        
    def add_documents(self, documents: List[Dict], source: str):
        """Add documents to vector store with metadata."""
        for doc in documents:
            doc.metadata["source"] = source
            doc.metadata["timestamp"] = format_timestamp()
            
        self.vectorstore.add_documents(documents)
        
    def similarity_search(self, query: str, k: int = 4) -> List[Dict]:
        """Perform similarity search with metadata filtering."""
        return self.vectorstore.similarity_search(
            query,
            k=k,
            search_distance=0.3
        )
        
    def get_retriever(self):
        """Get retriever for QA chains."""
        return self.vectorstore.as_retriever(
            search_kwargs={"k": 4}
        )
        
    def delete_document(self, doc_id: str) -> bool:
        """Delete a document from the vector store."""
        try:
            self.vectorstore.delete(doc_id)
            return True
        except Exception as e:
            logger.log_error(f"Error deleting document: {str(e)}")
            return False
            
    def delete_all(self):
        """Delete all documents from the vector store."""
        try:
            self.vectorstore.delete_all()
            return True
        except Exception as e:
            logger.log_error(f"Error deleting all documents: {str(e)}")
            return False

# Conversation Management
class ConversationManager:
    def __init__(self):
        self.memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True
        )
        
        self.qa_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=vector_store.get_retriever(),
            memory=self.memory,
            return_source_documents=True,
            verbose=True
        )
        
        self.processing_chain = self._create_processing_chain()
        
    def get_response(self, query: str) -> Dict:
        """Get response using conversation chain."""
        if not api_limiter.can_proceed():
            wait_time = api_limiter.wait_time()
            if wait_time > 0:
                time.sleep(wait_time)
                
        return self.qa_chain({"question": query})
        
    def process_text(self, text: str) -> Tuple[str, bool]:
        """Process text through improvement and verification chains."""
        return self.processing_chain.process(text)
        
    def clear_memory(self):
        """Clear conversation memory."""
        self.memory.clear()
        
    def _create_processing_chain(self):
        """Create text processing chain."""
        improve_prompt = PromptTemplate(
            input_variables=["text"],
            template="""
            Improve the clarity and readability of the following text while
            preserving all factual information:
            
            {text}
            
            Improved text:
            """
        )
        
        verify_prompt = PromptTemplate(
            input_variables=["original", "improved"],
            template="""
            Compare the original and improved texts for accuracy:
            
            Original: {original}
            Improved: {improved}
            
            Are there any factual inconsistencies? Respond with 'yes' or 'no':
            """
        )
        
        improve_chain = LLMChain(
            llm=llm,
            prompt=improve_prompt,
            verbose=True
        )
        
        verify_chain = LLMChain(
            llm=llm,
            prompt=verify_prompt,
            verbose=True
        )
        
        return SequentialChain(
            chains=[improve_chain, verify_chain],
            input_variables=["text"],
            output_variables=["improved_text", "verification"],
            verbose=True
        )

# Initialize managers
document_processor = DocumentProcessor()
vector_store = VectorStoreManager()
conversation_manager = ConversationManager()

# Progress Tracking
class ProgressTracker:
    def __init__(self, total_steps: int, description: str):
        self.progress_bar = st.progress(0)
        self.status_text = st.empty()
        self.total_steps = total_steps
        self.current_step = 0
        self.description = description
        
    def update(self, step: int = None, message: str = None):
        if step is not None:
            self.current_step = step
        if message is not None:
            self.status_text.text(f"{self.description}: {message}")
        self.progress_bar.progress(self.current_step / self.total_steps)
        
    def complete(self):
        self.progress_bar.progress(1.0)
        self.status_text.text(f"{self.description}: Completed!")
        time.sleep(1)
        self.progress_bar.empty()
        self.status_text.empty()

# Decorators
def safe_execute(func):
    """Decorator for safe execution with logging."""
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            logger.log_error(str(e), {'function': func.__name__, 'args': args, 'kwargs': kwargs})
            raise CognitextException(f"Error in {func.__name__}: {str(e)}")
    return wrapper

def process_with_memory_check(func):
    """Decorator to check memory before processing."""
    def wrapper(*args, **kwargs):
        estimated_size = sum(sys.getsizeof(arg) for arg in args)
        if not memory_manager.can_process(estimated_size):
            raise MemoryError("Insufficient memory for processing")
        return func(*args, **kwargs)
    return wrapper
import streamlit as st
from openai import OpenAI
import weaviate
from bs4 import BeautifulSoup
import requests
import PyPDF2
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import io
from docx import Document as DocxDocument
from fpdf import FPDF
from typing import List, Dict, Tuple
import os
from datetime import datetime
import pandas as pd
import time
import threading
from queue import Queue

# Initialize OpenAI and Weaviate clients
client_openai = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
client = weaviate.Client(
    url=st.secrets["WEAVIATE_URL"],
    auth_client_secret=weaviate.AuthApiKey(api_key=st.secrets["WEAVIATE_API_KEY"]),
    additional_headers={"X-OpenAI-Api-Key": st.secrets["OPENAI_API_KEY"]}
)
# Define a processing queue
if "processing_queue" not in st.session_state:
    st.session_state.processing_queue = Queue()
def transcribe_audio(audio_file: io.BytesIO) -> str:
    """Transcribe audio using OpenAI's Whisper API."""
    try:
        audio_bytes = audio_file.read()
        # Use OpenAI's Whisper API to transcribe
        transcription = client_openai.Audio.transcribe(
            model="whisper-1",  # Use the appropriate model here
            file=audio_bytes,
            response_format="text"
        )
        return transcription.get("text", "Transcription failed.")
    except Exception as e:
        st.error(f"Error transcribing audio: {str(e)}")
        return ""
def extract_text_from_webpage(url: str) -> str:
    """Extract text content from a webpage."""
    try:
        response = requests.get(url)
        response.raise_for_status()  # Check if the request was successful
        soup = BeautifulSoup(response.text, 'html.parser')

        # Extract text from paragraphs for readability
        text = ' '.join([p.get_text() for p in soup.find_all('p')])

        # Limit the content length to prevent overly long responses
        return text  # Limit to the first 10,000 characters if needed
    except requests.exceptions.RequestException as e:
        st.error(f"Failed to fetch content from the URL: {e}")
        return ""
    except Exception as e:
        st.error(f"An error occurred while processing the webpage: {e}")
        return ""

# Function to extract text from PDF using PyPDF2
def extract_text_from_pdf(pdf_file: io.BytesIO) -> str:
    """Extract text content from a PDF file."""
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    return text

# Function to extract text from DOCX
def extract_text_from_docx(docx_file: io.BytesIO) -> str:
    """Extract text content from a DOCX file."""
    doc = DocxDocument(docx_file)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)

# Function to create a DOCX file
def create_docx(text: str, filename: str) -> bytes:
    """Create a DOCX file from text."""
    doc = DocxDocument()
    doc.add_paragraph(text)
    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io.getvalue()

# Function to create a PDF file
def create_pdf(text: str, filename: str) -> bytes:
    """Create a PDF file from text."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    
    # Write each line to the PDF, handling encoding issues gracefully
    for line in text.split('\n'):
        try:
            # Encode each line as latin1, ignoring characters that cannot be encoded
            pdf.multi_cell(0, 10, line.encode('latin1', 'ignore').decode('latin1'))
        except Exception as e:
            st.error(f"Error writing line to PDF: {e}")
    
    # Create a BytesIO object to store the PDF data
    byte_io = io.BytesIO()
    pdf.output(dest='S').encode('latin1')
    byte_io.write(pdf.output(dest='S').encode('latin1'))
    byte_io.seek(0)
    return byte_io.getvalue()

# Function to allow download of the improved OCR document
def download_improved_ocr(content: str):
    """Provide options to download the improved OCR content as PDF or DOCX."""
    st.write("Download the improved OCR document:")
    docx_data = create_docx(content, "improved_ocr.docx")
    pdf_data = create_pdf(content, "improved_ocr.pdf")
    
    st.download_button(
        label="Download as DOCX",
        data=docx_data,
        file_name="improved_ocr.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    st.download_button(
        label="Download as PDF",
        data=pdf_data,
        file_name="improved_ocr.pdf",
        mime="application/pdf"
    )
def download_chatbot_document(response: str):
    """Provide options to download the chatbot's response as PDF or DOCX."""
    st.write("Download the generated document:")
    docx_data = create_docx(response, "chatbot_generated.docx")
    pdf_data = create_pdf(response, "chatbot_generated.pdf")
    
    st.download_button(
        label="Download as DOCX",
        data=docx_data,
        file_name="chatbot_generated.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    st.download_button(
        label="Download as PDF",
        data=pdf_data,
        file_name="chatbot_generated.pdf",
        mime="application/pdf"
    )
# Function to fetch all documents from Weaviate
def fetch_all_documents() -> List[Dict]:
    """Fetch all documents from the Weaviate database."""
    try:
        response = (
            client.query
            .get("Document", ["content", "source", "timestamp"])
            .with_additional(["id"])
            .do()
        )
        return response["data"]["Get"]["Document"]
    except Exception as e:
        st.error(f"Error fetching documents: {str(e)}")
        return []

# Function to delete a document from Weaviate
def delete_document_from_weaviate(doc_id: str) -> bool:
    """Delete a document from the Weaviate database using its ID."""
    try:
        client.data_object.delete(doc_id, class_name="Document")
        return True
    except Exception as e:
        st.error(f"Error deleting document: {str(e)}")
        return False

# Function for OCR using pytesseract
def extract_text_with_ocr(pdf_bytes: bytes) -> str:
    """Extract text from images in a PDF using OCR."""
    try:
        images = convert_pdf_to_images(pdf_bytes)
        extracted_text = ""
        for image in images:
            text = pytesseract.image_to_string(image)
            extracted_text += text + "\n"
        return extracted_text
    except Exception as e:
        st.error(f"Error during OCR processing: {str(e)}")
        return ""

def convert_pdf_to_images(pdf_bytes: bytes) -> List[Image.Image]:
    """Convert each page of a PDF to an image for OCR processing using pdf2image."""
    try:
        images = convert_from_bytes(pdf_bytes)
        return images
    except Exception as e:
        st.error(f"Error converting PDF to images: {str(e)}")
        return []

def improve_text_clarity(text: str, max_tokens_per_chunk: int = 3000) -> str:
    """Enhance the readability and clarity of the text using GPT in manageable chunks."""
    chunks = chunk_text(text, chunk_size=max_tokens_per_chunk)
    improved_chunks = []
    for i, chunk in enumerate(chunks):
        try:
            messages = [
                {"role": "system", "content": "You are an assistant tasked with improving the clarity and readability of a text without introducing any factual inaccuracies."},
                {"role": "user", "content": f"Improve the following text:\n\n{chunk}"}
            ]
            response = client_openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=messages,
                temperature=0.3,
                max_tokens=4000
            )
            improved_text = response.choices[0].message.content
            improved_chunks.append(improved_text)
        except Exception as e:
            st.error(f"Error improving chunk {i+1} with GPT: {str(e)}")
            improved_chunks.append(chunk)  # Add the original chunk if improvement fails
    
    return "\n\n".join(improved_chunks)

def verify_no_hallucinations(original_text: str, improved_text: str, max_tokens_per_chunk: int = 3000) -> bool:
    """Check if the improved text introduces any hallucinations or inaccuracies in manageable chunks."""
    original_chunks = chunk_text(original_text, chunk_size=max_tokens_per_chunk)
    improved_chunks = chunk_text(improved_text, chunk_size=max_tokens_per_chunk)
    
    for i, (orig_chunk, imp_chunk) in enumerate(zip(original_chunks, improved_chunks)):
        try:
            messages = [
                {"role": "system", "content": "You are a verification assistant. Compare two versions of a text and ensure the improved version does not introduce any hallucinations or inaccuracies."},
                {"role": "user", "content": f"Original text:\n{orig_chunk}\n\nImproved text:\n{imp_chunk}\n\nDoes the improved version introduce any inaccuracies? Respond with 'yes' or 'no'."}
            ]
            response = client_openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=messages,
                temperature=0,
                max_tokens=1000
            )
            verification = response.choices[0].message.content.strip().lower()
            if "yes" in verification:
                return False  # If any chunk introduces inaccuracies, fail the verification
        except Exception as e:
            st.error(f"Error verifying chunk {i+1} with GPT: {str(e)}")
            return False
    
    return True

def setup_weaviate():
    """Initialize Weaviate schema if it doesn't exist."""
    try:
        existing_schema = client.schema.get()
        schema_exists = any(class_obj["class"] == "Document" for class_obj in existing_schema["classes"]) if "classes" in existing_schema else False
        if not schema_exists:
            schema = {
                "class": "Document",
                "vectorizer": "text2vec-openai",
                "moduleConfig": {
                    "text2vec-openai": {
                        "model": "ada",
                        "modelVersion": "002",
                        "type": "text"
                    }
                },
                "properties": [
                    {"name": "content", "dataType": ["text"]},
                    {"name": "source", "dataType": ["string"]},
                    {"name": "timestamp", "dataType": ["date"]}
                ]
            }
            client.schema.create_class(schema)
            st.success("Schema created successfully")
    except Exception as e:
        st.error(f"Error setting up schema: {str(e)}")

def add_to_weaviate(content: str, source: str) -> bool:
    """Add content to Weaviate database."""
    chunks = chunk_text(content)
    try:
        for chunk in chunks:
            client.data_object.create(
                class_name="Document",
                data_object={
                    "content": chunk,
                    "source": source,
                    "timestamp": format_timestamp()
                }
            )
        return True
    except Exception as e:
        st.error(f"Error adding to Weaviate: {str(e)}")
        return False



def format_timestamp() -> str:
    """Format current timestamp in RFC3339 format."""
    return datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%S.%f')[:-3] + 'Z'
def process_materials():
    """Process all materials in the queue with progress tracking."""
    if st.session_state.processing_queue.empty():
        st.warning("No materials in the processing queue.")
        return

    # Create a progress bar
    total_files = st.session_state.processing_queue.qsize()
    progress_bar = st.progress(0)
    status_text = st.empty()
    files_processed = 0

    while not st.session_state.processing_queue.empty():
        try:
            # Get file info from queue (this also removes it from the queue)
            file_info = st.session_state.processing_queue.get()
            file_name = file_info["name"]
            file_bytes = file_info["bytes"]
            file_type = file_info["type"]
            use_ocr = file_info.get("use_ocr", False)

            # Update status
            status_text.text(f"Processing {file_name}...")

            # Process based on file type and OCR setting
            content = ""
            if use_ocr and file_type == "pdf":
                content = extract_text_with_ocr(file_bytes)
            elif file_type == "pdf":
                content = extract_text_from_pdf(io.BytesIO(file_bytes))
            elif file_type == "docx":
                content = extract_text_from_docx(io.BytesIO(file_bytes))
            else:  # txt files
                content = file_bytes.decode('utf-8')

            if not content.strip():
                st.warning(f"No content extracted from {file_name}")
                continue

            # Improve text clarity using GPT
            status_text.text(f"Improving text clarity for {file_name}...")
            improved_content = improve_text_clarity(content)

            # Add only the improved content to Weaviate database
            status_text.text(f"Adding improved version of {file_name} to database...")
            if add_to_weaviate(improved_content, file_name):
                st.success(f"Successfully processed and added improved version of {file_name}")
            else:
                st.error(f"Failed to add improved version of {file_name} to database")

            # Update progress
            files_processed += 1
            progress_bar.progress(files_processed / total_files)
            status_text.text(f"Processed {files_processed} of {total_files} files")

        except Exception as e:
            st.error(f"Unexpected error: {str(e)}")
            continue

    # Clear the progress bar and status text
    progress_bar.empty()
    status_text.empty()

    # Show final status
    if files_processed == total_files:
        st.success(f"Successfully processed all {files_processed} files")
    else:
        st.warning(f"Processed {files_processed} out of {total_files} files with some errors")

def delete_all_documents():
    """Delete all documents from the Weaviate database."""
    try:
        # Fetch all document IDs
        documents = fetch_all_documents()
        for doc in documents:
            doc_id = doc["_additional"]["id"]
            delete_document_from_weaviate(doc_id)
        st.success("All documents have been deleted from the library.")
    except Exception as e:
        st.error(f"Error deleting all documents: {str(e)}")
def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
    """Split text into chunks with overlap."""
    if not text:
        return []
        
    words = text.split()
    if not words:
        return []
        
    chunks = []
    i = 0
    while i < len(words):
        chunk_words = words[i:i + chunk_size]
        chunks.append(' '.join(chunk_words))
        i += (chunk_size - overlap)
    return chunks

def improve_text_clarity(text: str, max_tokens_per_chunk: int = 3000) -> str:
    """Enhance the readability and clarity of the text using GPT in manageable chunks."""
    if not text.strip():
        return text
        
    chunks = chunk_text(text, chunk_size=max_tokens_per_chunk)
    if not chunks:
        return text
        
    improved_chunks = []
    
    for i, chunk in enumerate(chunks):
        try:
            messages = [
                {"role": "system", "content": "You are an assistant tasked with improving the clarity and readability of text while preserving all factual information. Make the text more coherent and well-structured without changing its meaning."},
                {"role": "user", "content": f"Improve the following text:\n\n{chunk}"}
            ]
            
            response = client_openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=messages,
                temperature=0.3,
                max_tokens=4000
            )
            
            improved_text = response.choices[0].message.content
            if improved_text.strip():
                improved_chunks.append(improved_text)
            else:
                improved_chunks.append(chunk)  # Keep original if improvement is empty
                
        except Exception as e:
            st.error(f"Error improving chunk {i+1}: {str(e)}")
            improved_chunks.append(chunk)  # Keep original if improvement fails
    
    return "\n\n".join(improved_chunks)

def add_to_weaviate(content: str, source: str) -> bool:
    """Add content to Weaviate database with error handling and chunking."""
    if not content.strip():
        st.warning(f"No content to add for {source}")
        return False
        
    try:
        chunks = chunk_text(content)
        if not chunks:
            st.warning(f"No valid chunks created for {source}")
            return False
            
        for chunk in chunks:
            if chunk.strip():  # Only add non-empty chunks
                client.data_object.create(
                    class_name="Document",
                    data_object={
                        "content": chunk,
                        "source": source,
                        "timestamp": format_timestamp()
                    }
                )
        return True
        
    except Exception as e:
        st.error(f"Error adding content to Weaviate: {str(e)}")
        return False


def verify_document_in_weaviate(source_name: str) -> bool:
    """Verify if a specific document exists in Weaviate."""
    try:
        response = (
            client.query
            .get("Document", ["content", "source", "timestamp"])
            .with_where({
                "path": ["source"],
                "operator": "Equal",
                "valueString": source_name
            })
            .do()
        )
        
        if response and "data" in response and "Get" in response["data"]:
            documents = response["data"]["Get"]["Document"]
            return len(documents) > 0
        return False
    except Exception as e:
        st.error(f"Error verifying document: {str(e)}")
        return False

def list_all_documents() -> None:
    """List all documents in Weaviate with their content preview."""
    try:
        response = (
            client.query
            .get("Document", ["content", "source", "timestamp"])
            .with_additional(["id"])
            .do()
        )
        
        if response and "data" in response and "Get" in response["data"]:
            documents = response["data"]["Get"]["Document"]
            
            if not documents:
                st.warning("No documents found in the database.")
                return
                
            st.write("### Documents in Database:")
            for doc in documents:
                st.write(f"\n**Source:** {doc['source']}")
                st.write(f"**Preview:** {doc['content'][:200]}...")
                st.write("---")
        else:
            st.warning("No documents found in the database.")
            
    except Exception as e:
        st.error(f"Error listing documents: {str(e)}")

def improved_query_weaviate(query: str) -> List[Dict]:
    """Enhanced Weaviate query with better name matching and fuzzy search."""
    try:
        # First, check if we have any documents at all
        all_docs = (
            client.query
            .get("Document", ["content", "source"])
            .do()
        )
        
        if not all_docs.get("data", {}).get("Get", {}).get("Document", []):
            st.warning("No documents found in the database. Please upload documents first.")
            return []

        # 1. Try content-based search first
        content_response = (
            client.query
            .get("Document", ["content", "source", "timestamp"])
            .with_near_text({
                "concepts": [query],
                "certainty": 0.3  # Much lower threshold for better recall
            })
            .with_limit(limit)
            .do()
        )
        
        content_matches = content_response.get("data", {}).get("Get", {}).get("Document", [])

        # 2. Try searching within content using where filter
        words = query.lower().split()
        where_filter = {
            "operator": "Or",
            "operands": [
                {
                    "path": ["content"],
                    "operator": "Like",
                    "valueString": f"*{word}*"
                } for word in words
            ]
        }
        
        where_response = (
            client.query
            .get("Document", ["content", "source", "timestamp"])
            .with_where(where_filter)
            .with_limit(limit)
            .do()
        )
        
        where_matches = where_response.get("data", {}).get("Get", {}).get("Document", [])

        # Combine and deduplicate results
        all_matches = []
        seen_sources = set()
        
        # Process all matches
        for matches in [content_matches, where_matches]:
            for doc in matches:
                if doc['source'] not in seen_sources:
                    # Check if the content is relevant to the query
                    query_words = set(query.lower().split())
                    content_words = set(doc['content'].lower().split())
                    if query_words & content_words:  # If there's any word overlap
                        all_matches.append(doc)
                        seen_sources.add(doc['source'])

        # Debug output
        st.write("### Query Debug Information")
        st.write(f"Query: '{query}'")
        st.write(f"Total documents in database: {len(all_docs['data']['Get']['Document'])}")
        st.write(f"Semantic matches found: {len(content_matches)}")
        st.write(f"Keyword matches found: {len(where_matches)}")
        st.write(f"Combined unique relevant matches: {len(all_matches)}")
        
        if all_matches:
            st.write("\n### Matched Documents:")
            for idx, doc in enumerate(all_matches, 1):
                st.write(f"\n{idx}. Source: {doc['source']}")
                # Show more context around matched content
                content_preview = doc['content'][:300]
                st.write(f"Preview: {content_preview}...")
                
                # Highlight why this document was matched
                query_words = query.lower().split()
                matches_found = [word for word in query_words if word.lower() in content_preview.lower()]
                if matches_found:
                    st.write(f"Matched terms: {', '.join(matches_found)}")
        else:
            st.write("\nNo direct matches found. Showing all documents containing potentially relevant information:")
            for doc in all_docs['data']['Get']['Document']:
                if "Resume" in doc['source'] or any(word.lower() in doc['content'].lower() for word in query.split()):
                    st.write(f"\nSource: {doc['source']}")
                    st.write(f"Preview: {doc['content'][:300]}...")

        return all_matches if all_matches else [doc for doc in all_docs['data']['Get']['Document'] 
                                              if "Resume" in doc['source'] or 
                                              any(word.lower() in doc['content'].lower() for word in query.split())][:limit]
        
    except Exception as e:
        st.error(f"Error querying Weaviate: {str(e)}")
        return []

def get_chatgpt_response(prompt: str, context: List[Dict]) -> str:
    """Get response from ChatGPT using retrieved context with improved name handling."""
    if not context:
        return (
            "I don't have any relevant information about that in my document library. "
            "Please upload some relevant documents first."
        )
    
    # Format the context for the prompt
    context_text = "\n\n".join([
        f"Document '{doc['source']}':\n{doc['content']}"
        for doc in context
    ])
    
    try:
        messages = [
            {
                "role": "system",
                "content": (
                    "You are a helpful research assistant. Your task is to provide accurate, "
                    "comprehensive answers based on the provided documents. When answering:\n"
                    "1. Use only the information from the provided documents\n"
                    "2. If you find relevant information about a person, be thorough in describing their background\n"
                    "3. Always cite your sources using the document names\n"
                    "4. If the information is from a resume, present it in a professional manner\n"
                    "5. Structure your response for clarity"
                )
            },
            {
                "role": "user",
                "content": (
                    f"Based on the provided documents, please answer this question: {prompt}\n\n"
                    f"Documents for reference:\n{context_text}"
                )
            }
        ]
        
        response = client_openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages,
            temperature=0.3,
            max_tokens=2000
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        st.error(f"Error generating response: {str(e)}")
        return "I encountered an error while generating a response. Please try again."

def debug_chat_interface():
    st.header("Debug Chat Interface")
    
    # Document Database Status
    st.write("### Document Database Status")
    list_all_documents()
    
    # Test Query Section
    st.write("\n### Test Query")
    test_query = st.text_input("Enter a test query:")
    if test_query:
        st.write("### Performing test query...")
        docs = improved_query_weaviate(test_query)
        
        if docs:
            st.write("\n### Retrieved Documents for Response Generation:")
            # Create tabs for each document instead of expanders
            doc_tabs = st.tabs([f"Document {idx}: {doc['source']}" for idx, doc in enumerate(docs, 1)])
            
            for tab, doc in zip(doc_tabs, docs):
                with tab:
                    st.text_area("Content", doc['content'], height=200)
            
            st.write("\n### Generated Response:")
            response = get_chatgpt_response(test_query, docs)
            st.markdown(response)
        else:
            st.warning("No matching documents found for the query.")

def start_chat_interface():
    st.header("Chat with COGNITEXT")
    
    # Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Add debug toggle
    show_debug = st.checkbox("Show Debug Information")
    if show_debug:
        debug_chat_interface()
        st.divider()  # Add a visual separator

    # Display chat history
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Get user input
    if prompt := st.chat_input("What would you like to know about?"):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.spinner("Searching documents..."):
            relevant_docs = improved_query_weaviate(prompt)

        with st.chat_message("assistant"):
            if relevant_docs:
                response = get_chatgpt_response(prompt, relevant_docs)
                st.markdown(response)
                
                # Use columns for source display instead of expanders
                st.write("### Source Documents:")
                cols = st.columns(len(relevant_docs))
                for idx, (col, doc) in enumerate(zip(cols, relevant_docs), 1):
                    with col:
                        st.markdown(f"**Source {idx}: {doc['source']}**")
                        st.text_area(
                            f"Content",
                            doc['content'],
                            height=150,
                            key=f"source_{idx}_{doc['source']}"
                        )
            else:
                response = "I don't have any relevant information about that in my document library. Please upload some relevant documents first."
                st.markdown(response)

        st.session_state.messages.append({"role": "assistant", "content": response})
        download_chatbot_document(response)
def main():
    st.set_page_config(
        page_title="COGNITEXT",
        page_icon="📚",
        layout="wide"
    )

    # Initialize session state variables for modal states
    if "show_upload_modal" not in st.session_state:
        st.session_state.show_upload_modal = False
    if "show_url_modal" not in st.session_state:
        st.session_state.show_url_modal = False
    if "show_audio_modal" not in st.session_state:
        st.session_state.show_audio_modal = False
    setup_weaviate()
    # Define tabs for different sections
    tabs = st.tabs(["Main", "Document Library", "Chat"])
    # Main tab for chat and file handling
    with tabs[0]:
        st.markdown(
            "<h1 style='text-align: center; color: #3366cc;'>Welcome to COGNITEXT</h1>",
            unsafe_allow_html=True
        )
        st.markdown(
            "<p style='text-align: center; color: #666666;'>An AI-powered Assitant</p>",
            unsafe_allow_html=True
        )

        col1, col2, col3 = st.columns(3)

        with col1:
            st.image("upload_icon.png", width=175)
            if st.button("Upload Documents", key="upload_btn"):
                st.session_state.show_upload_modal = True
            
            if st.session_state.show_upload_modal:
               with st.expander("Upload Documents with OCR Option", expanded=True):
                use_ocr = st.radio("Use OCR for document processing?", options=["No", "Yes"], key="ocr_radio")
                uploaded_files = st.file_uploader(
                    "Upload Documents", 
                    type=["pdf", "docx", "txt"], 
                    accept_multiple_files=True,
                    key="doc_uploader"
                )

                if uploaded_files:
                    for uploaded_file in uploaded_files:
                        file_info = {
                            "name": uploaded_file.name,
                            "bytes": uploaded_file.getvalue(),
                            "type": uploaded_file.type.split('/')[-1],
                            "use_ocr": use_ocr == "Yes"
                        }
                        # Add file directly to processing function
                        st.session_state.processing_queue.put(file_info)
                    
                    # Call process_materials immediately after adding files
                    process_materials()

        with col2:
            st.image("url_icon.png", width=175)
            if st.button("Extract from URL", key="url_btn"):
                st.session_state.show_url_modal = True
            
            if st.session_state.show_url_modal:
                with st.expander("Extract Content from URL", expanded=True):
                    url = st.text_input("Enter webpage URL", key="url_input")
                    
                    # Add a close button
                    col2a, col2b = st.columns([4, 1])
                    with col2b:
                        if st.button("Close", key="close_url"):
                            st.session_state.show_url_modal = False
                            st.rerun()
                    
                    if url:
                        content = extract_text_from_webpage(url)
                        st.text_area("Extracted Content", content[:5000], height=300, key="url_content")
                        if st.button("Add to Weaviate", key="add_url"):
                            if add_to_weaviate(content, url):
                                st.success("Successfully added the extracted content to Weaviate.")
                                st.session_state.show_url_modal = False
                                st.rerun()

        with col3:
            st.image("audio_icon.png", width=175)
            if st.button("Upload Audio for Transcription", key="audio_btn"):
                st.session_state.show_audio_modal = True
            
            if st.session_state.show_audio_modal:
                with st.expander("Upload Audio Files for Transcription", expanded=True):
                    audio_files = st.file_uploader(
                        "Upload Audio Files", 
                        type=["mp3", "wav", "m4a"], 
                        accept_multiple_files=True,
                        key="audio_uploader"
                    )
                    
                    # Add a close button
                    col3a, col3b = st.columns([4, 1])
                    with col3b:
                        if st.button("Close", key="close_audio"):
                            st.session_state.show_audio_modal = False
                            st.rerun()
                    
                    if audio_files:
                        for audio_file in audio_files:
                            st.info(f"Transcribing {audio_file.name}...")
                            transcription = transcribe_audio(audio_file)
                            st.text_area(
                                f"Transcription of {audio_file.name}", 
                                transcription, 
                                height=300,
                                key=f"transcription_{audio_file.name}"
                            )
                            if st.button("Add to Weaviate", key=f"add_audio_{audio_file.name}"):
                                if add_to_weaviate(transcription, audio_file.name):
                                    st.success(f"Successfully added the transcription of {audio_file.name} to Weaviate.")
                                    st.session_state.show_audio_modal = False
                                    st.rerun()

        # Process materials button
        st.image("process_icon.png", width=175)
        if st.button("Process All Materials", key="process_btn"):
            process_materials()
            st.success("All materials processed successfully.")

        # Add a clear all button
        if st.button("Clear All Modals", key="clear_all"):
            st.session_state.show_upload_modal = False
            st.session_state.show_url_modal = False
            st.session_state.show_audio_modal = False
            st.rerun()

    # Document Library tab
    with tabs[1]:
        st.header("Document Library")
        documents = fetch_all_documents()
        if documents:
            df = pd.DataFrame(documents)
            df['timestamp'] = pd.to_datetime(df['timestamp'])
            st.dataframe(df[['source', 'timestamp']].rename(
                columns={'source': 'Document Source', 'timestamp': 'Uploaded At'}
            ))
            
            for idx, doc in enumerate(documents):
                doc_source = doc['source']
                doc_id = doc['_additional']['id']
                delete_key = f"delete_{doc_id}"
                if st.button(f"Delete '{doc_source}'", key=delete_key):
                    if delete_document_from_weaviate(doc_id):
                        st.success(f"Deleted '{doc_source}' successfully.")
                        time.sleep(1)  # Give user time to see the success message
                        st.rerun()
            # Add a "Delete All Files" button
        if st.button("Delete All Files"):
            delete_all_documents()

        else:
            st.write("No documents found in the library.")
    # Chat tab
    with tabs[2]:
        start_chat_interface()
# Run the main function
if __name__ == "__main__":
    main()