import streamlit as st
import weaviate
import openai
from typing import List, Dict
import os
from dotenv import load_dotenv
import pytesseract
from PIL import Image
import pdf2image
import fitz  # PyMuPDF
import cv2
import tempfile
from moviepy.editor import VideoFileClip
import base64
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import tempfile
from datetime import datetime, timedelta
import json
import pandas as pd
from collections import defaultdict
from io import BytesIO
import difflib

# Load environment variables
load_dotenv()

# Add this new class before your DocumentProcessor class:
class LegalReportGenerator:
    def __init__(self, openai_client):
        self.openai_client = openai_client
        self.report_templates = self._initialize_templates()
    
    def _initialize_templates(self) -> Dict[str, Dict]:
        """Initialize standard legal report templates"""
        return {
            "executive_summary": {
                "sections": [
                    "Overview",
                    "Key Findings",
                    "Risk Assessment",
                    "Recommendations"
                ],
                "style": "executive"
            },
            "detailed_analysis": {
                "sections": [
                    "Document Analysis",
                    "Party Analysis",
                    "Term Analysis",
                    "Risk Analysis",
                    "Compliance Analysis",
                    "Financial Analysis",
                    "Technical Requirements",
                    "Supporting Documentation"
                ],
                "style": "detailed"
            },
            "compliance_report": {
                "sections": [
                    "Compliance Summary",
                    "Regulatory Requirements",
                    "Compliance Status",
                    "Violations and Gaps",
                    "Remediation Plan",
                    "Monitoring Requirements"
                ],
                "style": "compliance"
            }
        }

    def generate_word_document(self, report_data: Dict, template_type: str = "detailed_analysis") -> bytes:
        """Generate a formatted Word document report"""
        doc = Document()
        
        # Set document properties
        doc.core_properties.author = "Legal RAG System"
        doc.core_properties.title = f"Legal Report - {report_data.get('document_id', 'Unnamed')}"
        
        # Add header
        self._add_header(doc, report_data)
        
        # Add table of contents
        self._add_table_of_contents(doc)
        
        # Add executive summary
        self._add_executive_summary(doc, report_data)
        
        # Add main content based on template
        template = self.report_templates.get(template_type, self.report_templates["detailed_analysis"])
        for section in template["sections"]:
            self._add_section(doc, section, report_data)
        
        # Add appendices
        self._add_appendices(doc, report_data)
        
        # Add footer
        self._add_footer(doc, report_data)
        
        # Save to bytes
        with tempfile.NamedTemporaryFile() as buffer:
            doc.save(buffer.name)
            buffer.seek(0)
            return buffer.read()

    def _add_header(self, doc: Document, report_data: Dict):
        """Add formatted header to document"""
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add logo if available
        # header_para.add_run().add_picture("logo.png", width=Inches(1.0))
        
        # Add company name and report title
        header_para.add_run("\nLegal Document Analysis Report").bold = True
        header_para.add_run(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    def _add_table_of_contents(self, doc: Document):
        """Add automated table of contents"""
        doc.add_paragraph("Table of Contents").style = 'Heading 1'
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

    def _add_executive_summary(self, doc: Document, report_data: Dict):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        
        summary = self._generate_executive_summary(report_data)
        for key, value in summary.items():
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(value)

    def _add_section(self, doc: Document, section_name: str, report_data: Dict):
        """Add a main content section"""
        doc.add_heading(section_name, level=1)
        
        # Get section content based on section name
        content = self._get_section_content(section_name, report_data)
        
        # Add content with appropriate formatting
        if isinstance(content, dict):
            for subsection, data in content.items():
                doc.add_heading(subsection, level=2)
                if isinstance(data, list):
                    for item in data:
                        doc.add_paragraph(item, style='List Bullet')
                else:
                    doc.add_paragraph(str(data))
        elif isinstance(content, list):
            for item in content:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph(str(content))

    def _add_appendices(self, doc: Document, report_data: Dict):
        """Add appendices with supporting documentation"""
        doc.add_heading('Appendices', level=1)
        
        # Add relevant appendices
        appendices = self._generate_appendices(report_data)
        for title, content in appendices.items():
            doc.add_heading(title, level=2)
            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(item, style='List Bullet')
            else:
                doc.add_paragraph(content)

    def _add_footer(self, doc: Document, report_data: Dict):
        """Add formatted footer to document"""
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_text = (
            f"Confidential - For Legal Review Only\n"
            f"Document ID: {report_data.get('document_id', 'Unnamed')}\n"
            f"Page "
        )
        footer_para.add_run(footer_text)
        
        # Add page numbers
        page_num = footer_para.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        page_num._r.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        page_num._r.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        page_num._r.append(fldChar)

    def _generate_executive_summary(self, report_data: Dict) -> Dict:
        """Generate executive summary content"""
        try:
            prompt = f"""
            Create an executive summary for this legal document analysis:
            {json.dumps(report_data, indent=2)}
            
            Include:
            1. Key Findings
            2. Risk Level
            3. Critical Issues
            4. Recommendations
            5. Next Steps
            
            Format as a structured JSON object with these sections as keys.
            """
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal document analyst creating executive summaries."},
                    {"role": "user", "content": prompt}
                ]
            )
            
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            return {
                "Error": f"Could not generate executive summary: {str(e)}",
                "Fallback Summary": "Please review the full report for details."
            }

    def _get_section_content(self, section_name: str, report_data: Dict) -> any:
        """Get content for a specific section"""
        if section_name in report_data:
            return report_data[section_name]
            
        # Generate content based on section name
        try:
            prompt = f"""
            Generate detailed content for the '{section_name}' section of this legal document analysis:
            {json.dumps(report_data, indent=2)}
            
            Provide structured, professional analysis appropriate for a legal document.
            Format the response as a JSON object with subsections and detailed points.
            """
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal document analyst creating detailed section content."},
                    {"role": "user", "content": prompt}
                ]
            )
            
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            return f"Content generation failed: {str(e)}"

    def _generate_appendices(self, report_data: Dict) -> Dict:
        """Generate appendices content"""
        appendices = {
            "Appendix A: Referenced Documents": self._get_referenced_documents(report_data),
            "Appendix B: Risk Analysis Details": self._get_detailed_risk_analysis(report_data),
            "Appendix C: Compliance Requirements": self._get_compliance_details(report_data),
            "Appendix D: Change History": self._get_change_history(report_data),
            "Appendix E: Supporting Materials": self._get_supporting_materials(report_data)
        }
        return appendices

    def _get_referenced_documents(self, report_data: Dict) -> List[str]:
        """Get list of referenced documents"""
        references = set()
        content = json.dumps(report_data)
        
        # Extract document references using GPT-4
        try:
            prompt = f"""
            Extract all document references from this content:
            {content}
            
            Include:
            1. Referenced contracts
            2. Supporting documents
            3. Related agreements
            4. Cited regulations
            5. Industry standards
            
            Return as a JSON array of strings.
            """
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are extracting document references from legal content."},
                    {"role": "user", "content": prompt}
                ]
            )
            
            return json.loads(response.choices[0].message.content).get("references", [])
        except Exception as e:
            return [f"Error extracting references: {str(e)}"]

    def _get_detailed_risk_analysis(self, report_data: Dict) -> Dict:
        """Get detailed risk analysis"""
        try:
            return self.risk_analyzer.analyze_risks(report_data) if hasattr(self, 'risk_analyzer') else {}
        except Exception as e:
            return {"error": f"Risk analysis failed: {str(e)}"}

    def _get_compliance_details(self, report_data: Dict) -> Dict:
        """Get detailed compliance information"""
        try:
            return self.compliance_checker.check_compliance(report_data) if hasattr(self, 'compliance_checker') else {}
        except Exception as e:
            return {"error": f"Compliance check failed: {str(e)}"}

    def _get_change_history(self, report_data: Dict) -> List[Dict]:
        """Get document change history"""
        return report_data.get("change_history", [])

    def _get_supporting_materials(self, report_data: Dict) -> List[str]:
        """Get list of supporting materials"""
        return report_data.get("supporting_materials", [])

    def export_to_docx(self, report_data: Dict, template_type: str = "detailed_analysis") -> bytes:
        """Export report to DOCX format"""
        return self.generate_word_document(report_data, template_type)

# Add this method to your RAGApplication class:
    def export_legal_report(self, contract_id: str, template_type: str = "detailed_analysis") -> bytes:
        """Generate and export a legal report in DOCX format"""
        try:
            # Gather all report data
            report_data = {
                "document_id": contract_id,
                "contract_info": self.contract_metadata.get(contract_id, {}),
                "legal_summary": self.get_formatted_legal_summary(contract_id),
                "due_diligence": self.get_due_diligence_report(contract_id),
                "compliance": self.get_compliance_report(contract_id),
                "audit_trail": self.get_contract_audit_trail(contract_id),
                "risk_assessment": self.risk_analyzer.analyze_risks(self.contract_metadata.get(contract_id, {})),
                "change_history": self._get_contract_activities(contract_id)
            }
            
            # Generate Word document
            report_generator = LegalReportGenerator(self.openai_client)
            return report_generator.export_to_docx(report_data, template_type)
            
        except Exception as e:
            raise Exception(f"Error generating legal report: {str(e)}")
class ContractAnalyzer:
    """Contract analysis and information extraction"""
    
    def __init__(self):
        self.contract_schema = {
            "parties": {
                "type": "array",
                "description": "List of all parties involved"
            },
            "effective_date": {
                "type": "string",
                "description": "Contract start date"
            },
            "expiration_date": {
                "type": "string",
                "description": "Contract end date"
            },
            "value": {
                "type": "object",
                "description": "Contract value and payment terms"
            },
            "key_clauses": {
                "type": "array",
                "description": "Important clauses and sections"
            },
            "obligations": {
                "type": "object",
                "description": "Party obligations"
            },
            "termination_conditions": {
                "type": "array",
                "description": "Termination terms"
            },
            "governing_law": {
                "type": "string",
                "description": "Governing jurisdiction"
            },
            "risk_factors": {
                "type": "array",
                "description": "Identified risks"
            }
        }
    
    def analyze_contract(self, content: str, openai_client) -> Dict[str, any]:
        """Analyze contract content using GPT-4"""
        try:
            prompt = f"""Analyze this contract as a legal expert and provide the following information in a structured format.
            
            Required Information:
            1. Parties: List all parties involved in the contract
            2. Effective Date: When the contract becomes effective
            3. Expiration Date: When the contract expires
            4. Contract Value: Details about monetary value and payment terms
            5. Key Clauses: List of important clauses and their content
            6. Obligations: Key obligations for each party
            7. Termination Conditions: Conditions for contract termination
            8. Governing Law: Jurisdiction governing the contract
            9. Risk Factors: Potential risks identified in the contract

            Format your response as follows:
            {
                "parties": ["Party 1", "Party 2"],
                "effective_date": "YYYY-MM-DD",
                "expiration_date": "YYYY-MM-DD",
                "value": {
                    "amount": "amount",
                    "currency": "currency",
                    "payment_terms": "terms"
                },
                "key_clauses": [
                    {"title": "clause title", "content": "clause content"}
                ],
                "obligations": {
                    "party1_name": ["obligation1", "obligation2"],
                    "party2_name": ["obligation1", "obligation2"]
                },
                "termination_conditions": ["condition1", "condition2"],
                "governing_law": "jurisdiction",
                "risk_factors": ["risk1", "risk2"]
            }

            Contract text:
            {content}

            Provide ONLY the JSON output without any additional text or explanation.
            """
            
            response = openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal document analyzer expert. Provide analysis in valid JSON format only."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7
            )
            
            # Parse the response text as JSON
            try:
                analysis = json.loads(response.choices[0].message.content)
                return analysis
            except json.JSONDecodeError:
                # If JSON parsing fails, try to extract JSON from the response
                content = response.choices[0].message.content
                # Find JSON-like structure between curly braces
                json_start = content.find('{')
                json_end = content.rfind('}') + 1
                if json_start >= 0 and json_end > json_start:
                    try:
                        analysis = json.loads(content[json_start:json_end])
                        return analysis
                    except json.JSONDecodeError:
                        return self._create_fallback_analysis(content)
                return self._create_fallback_analysis(content)
                
        except Exception as e:
            print(f"Contract analysis error: {str(e)}")
            return {
                "error": f"Error analyzing contract: {str(e)}",
                "parties": [],
                "effective_date": "Not found",
                "expiration_date": "Not found",
                "value": {"amount": "Unknown", "currency": "Unknown", "payment_terms": "Unknown"},
                "key_clauses": [],
                "obligations": {},
                "termination_conditions": [],
                "governing_law": "Not found",
                "risk_factors": []
            }

    def _create_fallback_analysis(self, content: str) -> Dict[str, any]:
        """Create a structured analysis from unstructured text"""
        try:
            # Extract information using basic text parsing
            analysis = {
                "parties": [],
                "effective_date": "Not found",
                "expiration_date": "Not found",
                "value": {
                    "amount": "Unknown",
                    "currency": "Unknown",
                    "payment_terms": "Unknown"
                },
                "key_clauses": [],
                "obligations": {},
                "termination_conditions": [],
                "governing_law": "Not found",
                "risk_factors": []
            }
            
            # Try to extract structured information from the text
            lines = content.split('\n')
            current_section = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Try to identify sections
                if "party" in line.lower() or "parties" in line.lower():
                    current_section = "parties"
                elif "effective date" in line.lower():
                    current_section = "effective_date"
                elif "expiration" in line.lower() or "termination date" in line.lower():
                    current_section = "expiration_date"
                elif "value" in line.lower() or "payment" in line.lower():
                    current_section = "value"
                elif "clause" in line.lower():
                    current_section = "key_clauses"
                elif "obligation" in line.lower():
                    current_section = "obligations"
                elif "terminate" in line.lower():
                    current_section = "termination_conditions"
                elif "govern" in line.lower() and "law" in line.lower():
                    current_section = "governing_law"
                elif "risk" in line.lower():
                    current_section = "risk_factors"
                
                # Process line based on current section
                if current_section:
                    if current_section == "parties":
                        if ":" in line:
                            party = line.split(":", 1)[1].strip()
                            if party and party not in analysis["parties"]:
                                analysis["parties"].append(party)
                    elif current_section in ["effective_date", "expiration_date"]:
                        if ":" in line:
                            date = line.split(":", 1)[1].strip()
                            analysis[current_section] = date
                    elif current_section == "key_clauses":
                        if ":" in line:
                            title, content = line.split(":", 1)
                            analysis["key_clauses"].append({
                                "title": title.strip(),
                                "content": content.strip()
                            })
                    elif current_section == "risk_factors":
                        if ":" in line:
                            risk = line.split(":", 1)[1].strip()
                            if risk and risk not in analysis["risk_factors"]:
                                analysis["risk_factors"].append(risk)
            
            return analysis
            
        except Exception as e:
            print(f"Error in fallback analysis: {str(e)}")
            return {
                "error": "Could not parse contract content",
                "parties": [],
                "effective_date": "Not found",
                "expiration_date": "Not found",
                "value": {"amount": "Unknown", "currency": "Unknown", "payment_terms": "Unknown"},
                "key_clauses": [],
                "obligations": {},
                "termination_conditions": [],
                "governing_law": "Not found",
                "risk_factors": []
            }
class ContractAlert:
    """Contract alert and reminder system"""
    def __init__(self):
        self.alerts = defaultdict(list)
        
    def add_alert(self, date_str: str, contract_id: str, alert_type: str, description: str):
        try:
            alert_date = datetime.fromisoformat(date_str)
            self.alerts[contract_id].append({
                "date": alert_date,
                "type": alert_type,
                "description": description,
                "status": "active"
            })
        except Exception as e:
            print(f"Error adding alert: {str(e)}")
    
    def get_upcoming_alerts(self, days_ahead: int = 30) -> List[Dict]:
        upcoming = []
        current_date = datetime.now()
        end_date = current_date + timedelta(days=days_ahead)
        
        for contract_id, contract_alerts in self.alerts.items():
            for alert in contract_alerts:
                if current_date <= alert["date"] <= end_date:
                    upcoming.append({
                        "contract_id": contract_id,
                        **alert
                    })
        
        return sorted(upcoming, key=lambda x: x["date"])

class ContractComparison:
    """Contract comparison and analysis tools"""
    def __init__(self, openai_client):
        self.openai_client = openai_client
    
    def compare_contracts(self, contract1: Dict, contract2: Dict) -> Dict:
        """Compare two contracts and identify differences"""
        comparison = {
            "parties_diff": self._compare_parties(contract1.get("parties", []), contract2.get("parties", [])),
            "dates_diff": self._compare_dates(contract1, contract2),
            "clauses_diff": self._compare_clauses(contract1.get("key_clauses", []), contract2.get("key_clauses", [])),
            "obligations_diff": self._compare_obligations(contract1.get("obligations", {}), contract2.get("obligations", {})),
            "risk_comparison": self._compare_risks(contract1, contract2)
        }
        
        return comparison
    
    def _compare_parties(self, parties1: List, parties2: List) -> Dict:
        """Compare parties between contracts"""
        return {
            "added": list(set(parties2) - set(parties1)),
            "removed": list(set(parties1) - set(parties2)),
            "common": list(set(parties1) & set(parties2))
        }
    
    def _compare_dates(self, contract1: Dict, contract2: Dict) -> Dict:
        """Compare important dates"""
        dates1 = {k: v for k, v in contract1.items() if 'date' in k.lower()}
        dates2 = {k: v for k, v in contract2.items() if 'date' in k.lower()}
        
        return {
            "different_dates": {
                k: {"contract1": dates1.get(k), "contract2": dates2.get(k)}
                for k in set(dates1.keys()) | set(dates2.keys())
                if dates1.get(k) != dates2.get(k)
            }
        }
    
    def _compare_clauses(self, clauses1: List, clauses2: List) -> Dict:
        """Compare contract clauses"""
        matcher = difflib.SequenceMatcher(None)
        similar_clauses = []
        different_clauses = []
        
        for c1 in clauses1:
            found_match = False
            for c2 in clauses2:
                matcher.set_seqs(c1, c2)
                if matcher.ratio() > 0.8:  # 80% similarity threshold
                    similar_clauses.append((c1, c2))
                    found_match = True
                    break
            if not found_match:
                different_clauses.append(c1)
        
        return {
            "similar": similar_clauses,
            "different": different_clauses,
            "unique_to_second": [c for c in clauses2 if c not in [x[1] for x in similar_clauses]]
        }

class RiskAnalyzer:
    """Contract risk analysis and assessment"""
    def __init__(self, openai_client):
        self.openai_client = openai_client
        
    def analyze_risks(self, contract_info: Dict) -> Dict:
        """Analyze contract risks and provide assessment"""
        try:
            risk_prompt = f"""Analyze the following contract information for potential risks:
            {json.dumps(contract_info, indent=2)}
            
            Provide a detailed risk assessment including:
            1. Risk level (High, Medium, Low)
            2. Risk categories
            3. Specific risk factors
            4. Mitigation suggestions
            
            Response in JSON format."""
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal risk assessment expert."},
                    {"role": "user", "content": risk_prompt}
                ],
                response_format={ "type": "json_object" }
            )
            
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            return {"error": str(e)}

class ContractExporter:
    """Contract export and reporting tools"""
    def generate_summary_report(self, contract_info: Dict) -> str:
        """Generate a detailed contract summary report"""
        report = f"""
        Contract Summary Report
        Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        
        1. Contract Parties:
        {self._format_parties(contract_info.get('parties', []))}
        
        2. Key Dates:
        - Effective Date: {contract_info.get('effective_date', 'Not specified')}
        - Expiration Date: {contract_info.get('expiration_date', 'Not specified')}
        
        3. Contract Value:
        {self._format_value(contract_info.get('value', {}))}
        
        4. Key Clauses:
        {self._format_clauses(contract_info.get('key_clauses', []))}
        
        5. Obligations:
        {self._format_obligations(contract_info.get('obligations', {}))}
        
        6. Risk Assessment:
        {self._format_risks(contract_info.get('risk_factors', []))}
        """
        return report
    
    def export_to_excel(self, contract_info: Dict) -> bytes:
        """Export contract information to Excel"""
        try:
            # Create different sheets for different aspects of the contract
            with pd.ExcelWriter() as writer:
                # Main Info
                pd.DataFrame([{
                    'Effective Date': contract_info.get('effective_date'),
                    'Expiration Date': contract_info.get('expiration_date'),
                    'Governing Law': contract_info.get('governing_law')
                }]).to_excel(writer, sheet_name='Main Info', index=False)
                
                # Parties
                pd.DataFrame(contract_info.get('parties', [])).to_excel(
                    writer, sheet_name='Parties', index=False
                )
                
                # Clauses
                pd.DataFrame(contract_info.get('key_clauses', [])).to_excel(
                    writer, sheet_name='Key Clauses', index=False
                )
                
                return writer.save()
        except Exception as e:
            raise Exception(f"Error exporting to Excel: {str(e)}")
    
    def _format_parties(self, parties: List) -> str:
        return "\n".join([f"- {party}" for party in parties])
    
    def _format_value(self, value: Dict) -> str:
        if isinstance(value, dict):
            return "\n".join([f"- {k}: {v}" for k, v in value.items()])
        return str(value)
    
    def _format_clauses(self, clauses: List) -> str:
        return "\n".join([f"- {clause}" for clause in clauses])
    
    def _format_obligations(self, obligations: Dict) -> str:
        return "\n".join([f"{party}:\n" + "\n".join([f"  - {obl}" for obl in obls])
                         for party, obls in obligations.items()])
    
    def _format_risks(self, risks: List) -> str:
        return "\n".join([f"- {risk}" for risk in risks])
           
class DocumentProcessor:
    def __init__(self, openai_client):
        self.openai_client = openai_client
        self.contract_analyzer = ContractAnalyzer()
        self.alert_system = ContractAlert()
        self.risk_analyzer = RiskAnalyzer(openai_client)
        self.contract_comparison = ContractComparison(openai_client)
        self.exporter = ContractExporter()
        self.processed_contracts = {}
    
    def process_text(self, content: str) -> str:
        """Process plain text content with enhanced contract analysis"""
        # Original text processing
        processed_text = content.strip()
        
        # Analyze for contract content
        contract_info = self.contract_analyzer.analyze_contract(processed_text, self.openai_client)
        if contract_info:
            # Store contract information
            contract_id = str(hash(processed_text))[:10]
            self.processed_contracts[contract_id] = contract_info
            
            # Set up alerts for important dates
            if 'effective_date' in contract_info:
                self.alert_system.add_alert(
                    contract_info['effective_date'],
                    contract_id,
                    'effective_date',
                    'Contract becomes effective'
                )
            if 'expiration_date' in contract_info:
                self.alert_system.add_alert(
                    contract_info['expiration_date'],
                    contract_id,
                    'expiration_date',
                    'Contract expires'
                )
            
            # Analyze risks
            risk_assessment = self.risk_analyzer.analyze_risks(contract_info)
            contract_info['risk_assessment'] = risk_assessment
            
            # Generate enhanced content with analysis
            enhanced_content = f"""
            ORIGINAL CONTENT:
            {processed_text}
            
            CONTRACT ANALYSIS:
            Parties Involved: {', '.join(contract_info.get('parties', []))}
            Effective Date: {contract_info.get('effective_date', 'Not specified')}
            Expiration Date: {contract_info.get('expiration_date', 'Not specified')}
            Contract Value: {contract_info.get('value', 'Not specified')}
            
            KEY CLAUSES:
            {self._format_clauses(contract_info.get('key_clauses', []))}
            
            RISK ASSESSMENT:
            {self._format_risk_assessment(risk_assessment)}
            """
            return enhanced_content
        
        return processed_text
    
    def _format_clauses(self, clauses: List) -> str:
        return '\n'.join([f'- {clause}' for clause in clauses])
    
    def _format_risk_assessment(self, assessment: Dict) -> str:
        if not assessment or 'error' in assessment:
            return "Risk assessment not available"
        
        return f"""
        Risk Level: {assessment.get('risk_level', 'Unknown')}
        Categories: {', '.join(assessment.get('risk_categories', []))}
        Key Factors: {', '.join(assessment.get('risk_factors', []))}
        """
    
    def get_contract_comparison(self, contract_id1: str, contract_id2: str) -> Dict:
        """Compare two processed contracts"""
        contract1 = self.processed_contracts.get(contract_id1)
        contract2 = self.processed_contracts.get(contract_id2)
        
        if not contract1 or not contract2:
            raise ValueError("One or both contract IDs not found")
        
        return self.contract_comparison.compare_contracts(contract1, contract2)
    
    def get_contract_report(self, contract_id: str, format: str = 'text') -> any:
        """Generate contract report in specified format"""
        contract_info = self.processed_contracts.get(contract_id)
        if not contract_info:
            raise ValueError("Contract ID not found")
        
        if format == 'text':
            return self.exporter.generate_summary_report(contract_info)
        elif format == 'excel':
            return self.exporter.export_to_excel(contract_info)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def process_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        text = []
        try:
            # Using PyMuPDF for text extraction
            doc = fitz.open(file_path)
            for page in doc:
                text.append(page.get_text())
            
            # If text extraction yields poor results, use OCR as fallback
            if not ''.join(text).strip():
                images = pdf2image.convert_from_path(file_path)
                for image in images:
                    text.append(pytesseract.image_to_string(image))
            
            return '\n'.join(text)
        except Exception as e:
            st.error(f"Error processing PDF: {str(e)}")
            return ""

    def process_image(self, file_path: str) -> Dict[str, str]:
        """Process image files using OCR and image analysis"""
        try:
            # Perform OCR
            image = Image.open(file_path)
            ocr_text = pytesseract.image_to_string(image)
            
            # Get image description using OpenAI Vision
            with open(file_path, "rb") as image_file:
                image_data = base64.b64encode(image_file.read()).decode('utf-8')
                
            response = self.openai_client.chat.completions.create(
                model="gpt-4-vision-preview",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "Describe this image in detail."},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{image_data}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=500
            )
            
            return {
                "ocr_text": ocr_text,
                "description": response.choices[0].message.content
            }
        except Exception as e:
            st.error(f"Error processing image: {str(e)}")
            return {"ocr_text": "", "description": ""}

    def process_audio(self, file_path: str) -> str:
        """Process audio files using OpenAI Whisper API"""
        try:
            with open(file_path, "rb") as audio_file:
                response = self.openai_client.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio_file
                )
                return response.text
        except Exception as e:
            st.error(f"Error processing audio: {str(e)}")
            return ""

    def process_video(self, file_path: str) -> Dict[str, str]:
        """Process video files - extract audio for transcription and frames for analysis"""
        try:
            # Extract audio and transcribe
            video = VideoFileClip(file_path)
            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_audio:
                video.audio.write_audiofile(temp_audio.name)
                transcription = self.process_audio(temp_audio.name)
                os.unlink(temp_audio.name)
            
            # Extract frames and analyze
            cap = cv2.VideoCapture(file_path)
            frames = []
            frame_descriptions = []
            
            frame_count = 0
            total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            frame_interval = max(1, total_frames // 5)  # Get 5 evenly spaced frames
            
            while cap.isOpened():
                ret, frame = cap.read()
                if not ret:
                    break
                
                if frame_count % frame_interval == 0 and len(frames) < 5:
                    frames.append(frame)
                
                frame_count += 1
            
            cap.release()
            
            # Analyze key frames
            for frame in frames:
                with tempfile.NamedTemporaryFile(suffix=".jpg") as temp_frame:
                    cv2.imwrite(temp_frame.name, frame)
                    frame_info = self.process_image(temp_frame.name)
                    frame_descriptions.append(frame_info["description"])
            
            return {
                "transcription": transcription,
                "frame_descriptions": "\n".join(frame_descriptions)
            }
        except Exception as e:
            st.error(f"Error processing video: {str(e)}")
            return {"transcription": "", "frame_descriptions": ""}

    def get_file_type(self, filename: str) -> str:
        """Determine file type based on extension"""
        ext = Path(filename).suffix.lower()
        if ext in ['.txt']:
            return 'text/plain'
        elif ext in ['.pdf']:
            return 'application/pdf'
        elif ext in ['.png', '.jpg', '.jpeg']:
            return 'image/' + ext[1:]
        elif ext in ['.mp3', '.wav']:
            return 'audio/' + ext[1:]
        elif ext in ['.mp4', '.avi']:
            return 'video/' + ext[1:]
        else:
            return 'application/octet-stream'

class RAGApplication:
    def __init__(self, openai_api_key: str, weaviate_url: str, weaviate_api_key: str):
        # Initialize OpenAI client
        self.openai_client = openai.Client(api_key=openai_api_key)
        
        # Initialize document processor
        self.processor = DocumentProcessor(self.openai_client)
        
        # Initialize Weaviate client with authentication
        auth_config = weaviate.auth.AuthApiKey(api_key=weaviate_api_key)
        
        self.weaviate_client = weaviate.Client(
            url=weaviate_url,
            auth_client_secret=auth_config,
            additional_headers={
                "X-OpenAI-Api-Key": openai_api_key
            }
        )
        
        # Class name for Weaviate
        self.class_name = "Document"
        
        # Initialize contract tracking
        self.contract_alerts = {}
        self.contract_metadata = {}
        self.contract_templates = {}
        
        # Create schema
        self._create_schema()
        
    def _create_schema(self):
        """Create enhanced schema with legal metadata support"""
        schema = {
            "classes": [{
                "class": self.class_name,
                "vectorizer": "text2vec-openai",
                "properties": [
                    {
                        "name": "content",
                        "dataType": ["text"],
                        "vectorizer": "text2vec-openai"
                    },
                    {
                        "name": "title",
                        "dataType": ["string"],
                        "vectorizer": "text2vec-openai"
                    },
                    {
                        "name": "file_type",
                        "dataType": ["string"]
                    },
                    {
                        "name": "metadata",
                        "dataType": ["text"]
                    },
                    {
                        "name": "contract_type",
                        "dataType": ["string"]
                    },
                    {
                        "name": "parties",
                        "dataType": ["string[]"]
                    },
                    {
                        "name": "effective_date",
                        "dataType": ["date"]
                    },
                    {
                        "name": "expiration_date",
                        "dataType": ["date"]
                    },
                    {
                        "name": "risk_level",
                        "dataType": ["string"]
                    },
                    {
                        "name": "legal_category",
                        "dataType": ["string"]
                    }
                ]
            }]
        }
        
        try:
            existing_schema = self.weaviate_client.schema.get()
            existing_classes = [c['class'] for c in existing_schema['classes']] if existing_schema.get('classes') else []
            
            if self.class_name not in existing_classes:
                self.weaviate_client.schema.create_class(schema['classes'][0])
        except Exception as e:
            print(f"Error with schema: {e}")

    def process_file(self, file, filename: str) -> Dict[str, str]:
        """Process file with enhanced legal document analysis"""
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as temp_file:
            temp_file.write(file.getvalue())
            file_path = temp_file.name
        
        try:
            file_type = self.processor.get_file_type(filename)
            
            # Process based on file type
            if file_type.startswith('image/'):
                result = self.processor.process_image(file_path)
                content = f"OCR Text: {result['ocr_text']}\nImage Description: {result['description']}"
            elif file_type.startswith('audio/'):
                content = self.processor.process_audio(file_path)
            elif file_type.startswith('video/'):
                result = self.processor.process_video(file_path)
                content = f"Transcription: {result['transcription']}\nVideo Description: {result['frame_descriptions']}"
            elif file_type == 'application/pdf':
                content = self.processor.process_pdf(file_path)
            else:
                content = self.processor.process_text(file.getvalue().decode())
            
            # Analyze for contract content
            contract_info = None
            if file_type in ['text/plain', 'application/pdf']:
                try:
                    contract_info = self.processor.contract_analyzer.analyze_contract(content, self.openai_client)
                    if contract_info:
                        # Generate contract ID
                        contract_id = str(hash(content))[:10]
                        
                        # Store contract metadata
                        self.contract_metadata[contract_id] = {
                            "filename": filename,
                            "analysis_date": datetime.now().isoformat(),
                            "contract_info": contract_info
                        }
                        
                        # Set up alerts
                        self.contract_alerts[contract_id] = {
                            "effective_date": contract_info.get("effective_date"),
                            "expiration_date": contract_info.get("expiration_date"),
                            "review_dates": []
                        }
                        
                        # Analyze risks
                        risk_assessment = self.processor.risk_analyzer.analyze_risks(contract_info)
                        
                        # Enhanced metadata for legal documents
                        metadata = {
                            "contract_id": contract_id,
                            "document_type": "contract",
                            "parties": contract_info.get("parties", []),
                            "effective_date": contract_info.get("effective_date"),
                            "expiration_date": contract_info.get("expiration_date"),
                            "risk_level": risk_assessment.get("risk_level"),
                            "risk_factors": risk_assessment.get("risk_factors", []),
                            "legal_category": self._determine_legal_category(contract_info),
                            "processing_date": datetime.now().isoformat()
                        }
                except Exception as e:
                    print(f"Contract analysis error: {str(e)}")
                    metadata = {
                        "filename": filename,
                        "file_type": file_type,
                        "processing_date": datetime.now().isoformat()
                    }
            else:
                metadata = {
                    "filename": filename,
                    "file_type": file_type,
                    "processing_date": datetime.now().isoformat()
                }
            
            return {
                "content": content,
                "file_type": file_type,
                "metadata": metadata,
                "contract_info": contract_info
            }
        finally:
            os.unlink(file_path)

    def add_document(self, title: str, file_data: Dict[str, str]) -> bool:
        """Add document with enhanced metadata"""
        try:
            # Split content into chunks
            content_chunks = self._chunk_text(file_data["content"])
            st.write(f"Splitting document into {len(content_chunks)} chunks")
            
            # Add each chunk as a separate document
            for i, chunk in enumerate(content_chunks):
                chunk_title = f"{title} (Part {i+1}/{len(content_chunks)})" if len(content_chunks) > 1 else title
                
                properties = {
                    "title": chunk_title,
                    "content": chunk,
                    "file_type": file_data["file_type"],
                    "metadata": file_data["metadata"]
                }
                
                # Add legal-specific properties if available
                if "contract_info" in file_data and file_data["contract_info"]:
                    contract_info = file_data["contract_info"]
                    properties.update({
                        "contract_type": contract_info.get("contract_type", "unknown"),
                        "parties": contract_info.get("parties", []),
                        "effective_date": contract_info.get("effective_date"),
                        "expiration_date": contract_info.get("expiration_date"),
                        "risk_level": contract_info.get("risk_level", "unknown"),
                        "legal_category": contract_info.get("legal_category", "unknown")
                    })
                
                result = self.weaviate_client.data_object.create(
                    class_name=self.class_name,
                    data_object=properties
                )
                
                if not result:
                    print(f"Failed to add chunk {i+1}")
                    return False
                
                st.write(f"Added chunk {i+1}/{len(content_chunks)}")
            
            return True
            
        except Exception as e:
            print(f"Error adding document: {e}")
            st.error(f"Error adding document: {e}")
            return False

    def search_documents(self, query: str, filters: Dict = None, limit: int = 3) -> List[dict]:
        """Enhanced search with legal document filtering"""
        try:
            # Build search query
            search_query = (
                self.weaviate_client.query
                .get(self.class_name, ["title", "content", "file_type", "metadata", 
                                     "contract_type", "parties", "effective_date", 
                                     "expiration_date", "risk_level", "legal_category"])
                .with_near_text({"concepts": [query]})
            )
            
            # Add filters if provided
            if filters:
                filter_conditions = []
                if filters.get("contract_type"):
                    filter_conditions.append({
                        "path": ["contract_type"],
                        "operator": "Equal",
                        "valueString": filters["contract_type"]
                    })
                if filters.get("risk_level"):
                    filter_conditions.append({
                        "path": ["risk_level"],
                        "operator": "Equal",
                        "valueString": filters["risk_level"]
                    })
                if filters.get("legal_category"):
                    filter_conditions.append({
                        "path": ["legal_category"],
                        "operator": "Equal",
                        "valueString": filters["legal_category"]
                    })
                if filters.get("date_range"):
                    date_range = filters["date_range"]
                    filter_conditions.extend([
                        {
                            "path": ["effective_date"],
                            "operator": "GreaterThanEqual",
                            "valueDate": date_range["start"]
                        },
                        {
                            "path": ["effective_date"],
                            "operator": "LessThanEqual",
                            "valueDate": date_range["end"]
                        }
                    ])
                
                if filter_conditions:
                    search_query = search_query.with_where({
                        "operator": "And",
                        "operands": filter_conditions
                    })
            
            # Execute search
            result = search_query.with_limit(limit).do()
            
            if result and "data" in result and "Get" in result["data"] and self.class_name in result["data"]["Get"]:
                return result["data"]["Get"][self.class_name]
            return []
            
        except Exception as e:
            print(f"Error searching documents: {e}")
            return []

    def generate_response(self, query: str, context_docs: List[dict]) -> str:
        """Generate enhanced response with legal context awareness"""
        try:
            # Combine context with enhanced legal awareness
            context_parts = []
            total_length = 0
            max_context_length = 4000
            
            for doc in context_docs:
                doc_text = f"""
                Title: {doc['title']}
                Type: {doc.get('legal_category', doc['file_type'])}
                Content: {doc['content']}
                
                Additional Context:
                - Contract Type: {doc.get('contract_type', 'N/A')}
                - Risk Level: {doc.get('risk_level', 'N/A')}
                - Effective Date: {doc.get('effective_date', 'N/A')}
                - Parties Involved: {', '.join(doc.get('parties', []))}
                """
                
                if total_length + len(doc_text) > max_context_length:
                    break
                    
                context_parts.append(doc_text)
                total_length += len(doc_text)
            
            context = "\n\n".join(context_parts)
            
            # Create prompt with legal context awareness
            prompt = f"""As a legal expert, analyze the following context and question.
            Provide a detailed answer considering legal implications and risks.
            
            Context:
            {context}

            Question:
            {query}

            Provide a structured response including:
            1. Direct answer to the question
            2. Legal considerations
            3. Risk factors (if any)
            4. Recommendations (if applicable)
            """
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal expert assistant providing accurate and detailed legal information."},
                    {"role": "user", "content": prompt}
                ]
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            return f"Error generating response: {e}"

    def get_contract_alerts(self, days_ahead: int = 30) -> List[Dict]:
        """Get upcoming contract alerts"""
        return self.processor.alert_system.get_upcoming_alerts(days_ahead)

    def compare_contracts(self, contract_id1: str, contract_id2: str) -> Dict:
        """Compare two contracts"""
        return self.processor.get_contract_comparison(contract_id1, contract_id2)

    def generate_contract_report(self, contract_id: str, format: str = 'text') -> any:
        """Generate contract report"""
        return self.processor.get_contract_report(contract_id, format)

    def _chunk_text(self, text: str, chunk_size: int = 2000) -> List[str]:
        """Split text into smaller chunks"""
        words = text.split()
        chunks = []
        current_chunk = []
        current_length = 0
        
        for word in words:
            current_length += len(word) + 1
            if current_length > chunk_size:
                chunks.append(' '.join(current_chunk))
                current_chunk = [word]
                current_length = len(word)
            else:
                current_chunk.append(word)
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks

    def _determine_legal_category(self, contract_info: Dict) -> str:
        """Determine the legal category of a contract"""
        # Add logic to categorize contracts based on content
        keywords = {
            "employment": ["employment", "worker", "salary", "compensation"],
            "service": ["services", "contractor", "scope of work"],
            "lease": ["lease", "rent", "property", "premises"],
            "sales": ["purchase", "sale", "goods", "delivery"],
            "nda": ["confidential", "non-disclosure", "proprietary"],
            "licensing": ["license", "intellectual property", "patent"]
        }
        
        content = json.dumps(contract_info).lower()
        
        for category, terms in keywords.items():
            if any(term in content for term in terms):
                return category
                
        return "other"
    def generate_legal_summary(self, contract_info: Dict) -> Dict:
        """Generate a legally formatted document summary"""
        try:
            prompt = f"""
            As a legal document specialist, create a formal legal summary of this contract in JSON format:
            {json.dumps(contract_info, indent=2)}
            
            Structure the response as follows:
            {{
                "document_identification": {{ "id": "", "type": "", "date": "" }},
                "parties_and_roles": [ {{ "party": "", "role": "" }} ],
                "key_terms": [ {{ "term": "", "definition": "" }} ],
                "material_obligations": [ {{ "party": "", "obligations": [] }} ],
                "critical_dates": [ {{ "event": "", "date": "" }} ],
                "financial_terms": {{ "value": "", "payment_terms": "" }},
                "termination_conditions": [],
                "governing_law": "",
                "special_provisions": []
            }}
            
            Provide ONLY the JSON output without any additional text.
            """
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal document specialist creating formal summaries in JSON format."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7
            )
            
            try:
                return json.loads(response.choices[0].message.content)
            except json.JSONDecodeError:
                # Extract JSON from response if needed
                content = response.choices[0].message.content
                json_start = content.find('{')
                json_end = content.rfind('}') + 1
                if json_start >= 0 and json_end > json_start:
                    return json.loads(content[json_start:json_end])
                return {
                    "error": "Could not parse response",
                    "raw_content": content
                }
                
        except Exception as e:
            return {
                "error": f"Error generating legal summary: {str(e)}",
                "document_id": contract_info.get("contract_id")
            }
    def extract_key_dates(self, contract_id: str) -> List[Dict]:
        """Extract all important dates from a contract"""
        try:
            contract_info = self.contract_metadata.get(contract_id)
            if not contract_info:
                return []
            
            dates = []
            if 'effective_date' in contract_info:
                dates.append({
                    "type": "effective_date",
                    "date": contract_info['effective_date'],
                    "description": "Contract becomes effective"
                })
            
            if 'expiration_date' in contract_info:
                dates.append({
                    "type": "expiration_date",
                    "date": contract_info['expiration_date'],
                    "description": "Contract expires"
                })
            
            # Extract additional dates from content
            content = json.dumps(contract_info)
            date_extraction_prompt = f"""
            Extract all important dates mentioned in this contract:
            {content}
            
            Include payment dates, review dates, notification deadlines, etc.
            Format as a list of JSON objects with date, type, and description.
            """
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal date extraction specialist."},
                    {"role": "user", "content": date_extraction_prompt}
                ],
                response_format={ "type": "json_object" }
            )
            
            additional_dates = json.loads(response.choices[0].message.content).get("dates", [])
            dates.extend(additional_dates)
            
            return sorted(dates, key=lambda x: x["date"])
        except Exception as e:
            return []

    def analyze_compliance(self, contract_id: str) -> Dict:
        """Analyze contract compliance requirements"""
        try:
            contract_info = self.contract_metadata.get(contract_id)
            if not contract_info:
                raise ValueError("Contract not found")
            
            compliance_prompt = f"""
            Analyze this contract for compliance requirements:
            {json.dumps(contract_info, indent=2)}
            
            Provide:
            1. Required compliance measures
            2. Reporting requirements
            3. Audit requirements
            4. Regulatory considerations
            5. Compliance risk assessment
            
            Format as a detailed JSON object.
            """
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a legal compliance expert."},
                    {"role": "user", "content": compliance_prompt}
                ],
                response_format={ "type": "json_object" }
            )
            
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            return {"error": f"Error analyzing compliance: {str(e)}"}

    def find_similar_contracts(self, contract_id: str, min_similarity: float = 0.7) -> List[Dict]:
        """Find similar contracts in the database"""
        try:
            contract_info = self.contract_metadata.get(contract_id)
            if not contract_info:
                return []
            
            # Create embedding for the contract
            contract_text = json.dumps(contract_info)
            
            # Search for similar contracts
            result = (
                self.weaviate_client.query
                .get(self.class_name, ["title", "content", "metadata", "contract_type"])
                .with_near_text({"concepts": [contract_text]})
                .with_limit(5)
                .do()
            )
            
            similar_contracts = []
            if result and "data" in result and "Get" in result["data"]:
                for doc in result["data"]["Get"][self.class_name]:
                    # Calculate similarity score
                    similarity = self._calculate_similarity(contract_text, doc["content"])
                    if similarity >= min_similarity:
                        similar_contracts.append({
                            "title": doc["title"],
                            "similarity": similarity,
                            "contract_type": doc.get("contract_type", "unknown"),
                            "metadata": doc["metadata"]
                        })
            
            return sorted(similar_contracts, key=lambda x: x["similarity"], reverse=True)
        except Exception as e:
            return []

    def get_contract_statistics(self) -> Dict:
        """Generate statistics about stored contracts"""
        try:
            stats = {
                "total_contracts": len(self.contract_metadata),
                "by_type": defaultdict(int),
                "by_risk_level": defaultdict(int),
                "by_status": defaultdict(int),
                "average_processing_time": 0,
                "expiring_soon": 0
            }
            
            processing_times = []
            current_date = datetime.now()
            
            for contract_id, metadata in self.contract_metadata.items():
                # Count by type
                contract_type = metadata.get("contract_info", {}).get("contract_type", "unknown")
                stats["by_type"][contract_type] += 1
                
                # Count by risk level
                risk_level = metadata.get("contract_info", {}).get("risk_level", "unknown")
                stats["by_risk_level"][risk_level] += 1
                
                # Check status and expiration
                if "expiration_date" in metadata:
                    exp_date = datetime.fromisoformat(metadata["expiration_date"])
                    if exp_date < current_date:
                        stats["by_status"]["expired"] += 1
                    elif exp_date < current_date + timedelta(days=30):
                        stats["by_status"]["expiring_soon"] += 1
                        stats["expiring_soon"] += 1
                    else:
                        stats["by_status"]["active"] += 1
                
                # Calculate processing time if available
                if "analysis_date" in metadata:
                    processing_time = datetime.fromisoformat(metadata["analysis_date"]) - \
                                   datetime.fromisoformat(metadata["processing_date"])
                    processing_times.append(processing_time.total_seconds())
            
            if processing_times:
                stats["average_processing_time"] = sum(processing_times) / len(processing_times)
            
            return stats
        except Exception as e:
            return {"error": f"Error generating statistics: {str(e)}"}

    def export_contract_database(self, format: str = 'excel') -> any:
        """Export the entire contract database"""
        try:
            data = []
            for contract_id, metadata in self.contract_metadata.items():
                contract_info = metadata.get("contract_info", {})
                data.append({
                    "contract_id": contract_id,
                    "filename": metadata.get("filename"),
                    "contract_type": contract_info.get("contract_type"),
                    "parties": ", ".join(contract_info.get("parties", [])),
                    "effective_date": contract_info.get("effective_date"),
                    "expiration_date": contract_info.get("expiration_date"),
                    "risk_level": contract_info.get("risk_level"),
                    "processing_date": metadata.get("processing_date")
                })
            
            if format == 'excel':
                output = BytesIO()
                df = pd.DataFrame(data)
                df.to_excel(output, index=False)
                return output.getvalue()
            elif format == 'json':
                return json.dumps(data, indent=2)
            else:
                raise ValueError(f"Unsupported export format: {format}")
        except Exception as e:
            raise Exception(f"Error exporting database: {str(e)}")

    def _calculate_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity between two texts"""
        try:
            # Use OpenAI embeddings for similarity calculation
            response1 = self.openai_client.embeddings.create(
                model="text-embedding-ada-002",
                input=text1
            )
            response2 = self.openai_client.embeddings.create(
                model="text-embedding-ada-002",
                input=text2
            )
            
            embedding1 = response1.data[0].embedding
            embedding2 = response2.data[0].embedding
            
            # Calculate cosine similarity
            dot_product = sum(a * b for a, b in zip(embedding1, embedding2))
            norm1 = sum(a * a for a in embedding1) ** 0.5
            norm2 = sum(b * b for b in embedding2) ** 0.5
            
            return dot_product / (norm1 * norm2)
        except Exception as e:
            print(f"Error calculating similarity: {str(e)}")
            return 0.0

    def backup_database(self, backup_path: str) -> bool:
        """Create a backup of the contract database"""
        try:
            backup_data = {
                "contract_metadata": self.contract_metadata,
                "contract_alerts": self.contract_alerts,
                "contract_templates": self.contract_templates,
                "backup_date": datetime.now().isoformat()
            }
            
            with open(backup_path, 'w') as f:
                json.dump(backup_data, f, indent=2)
            
            return True
        except Exception as e:
            print(f"Error creating backup: {str(e)}")
            return False

    def restore_database(self, backup_path: str) -> bool:
        """Restore the contract database from a backup"""
        try:
            with open(backup_path, 'r') as f:
                backup_data = json.load(f)
            
            self.contract_metadata = backup_data["contract_metadata"]
            self.contract_alerts = backup_data["contract_alerts"]
            self.contract_templates = backup_data["contract_templates"]
            
            return True
        except Exception as e:
            print(f"Error restoring backup: {str(e)}")
            return False

def main():
    st.title("📚 Legal Document Analysis System")
    
    # Load environment variables
    load_dotenv()
    
    # Get API keys from environment variables
    openai_api_key = os.getenv("OPENAI_API_KEY")
    weaviate_url = os.getenv("WEAVIATE_URL")
    weaviate_api_key = os.getenv("WEAVIATE_API_KEY")
    
    # Verify that all required environment variables are set
    if not all([openai_api_key, weaviate_url, weaviate_api_key]):
        st.error("Missing required environment variables. Please check your .env file.")
        st.info("Required variables: OPENAI_API_KEY, WEAVIATE_URL, WEAVIATE_API_KEY")
        return
    
    try:
        # Initialize RAG application
        rag_app = RAGApplication(openai_api_key, weaviate_url, weaviate_api_key)
        st.success("Successfully connected to services!")
        
        # Create main navigation
        menu = ["Document Management", "Search & Analysis", "Reports & Export", "Settings"]
        choice = st.sidebar.selectbox("Navigation", menu)
        
        if choice == "Document Management":
            show_document_management(rag_app)
        elif choice == "Search & Analysis":
            show_search_analysis(rag_app)
        elif choice == "Reports & Export":
            show_reports_export(rag_app)
        else:
            show_settings(rag_app)
            
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")
        st.info("Please check your API keys and connections.")
def show_document_management(rag_app):
    """Document management interface"""
    st.header("📁 Document Management")
    
    # Create tabs for different document operations
    doc_tabs = st.tabs(["Upload", "Manage", "Templates"])
    
    # Upload Tab
    with doc_tabs[0]:
        st.subheader("Upload Documents")
        uploaded_file = st.file_uploader(
            "Choose a file (Text, PDF, Image, Audio, or Video)", 
            type=["txt", "pdf", "png", "jpg", "jpeg", "mp3", "wav", "mp4", "avi"]
        )
        
        if uploaded_file:
            st.write(f"Selected file: {uploaded_file.name}")
            file_stats = f"File size: {uploaded_file.size / 1024:.2f} KB"
            st.write(file_stats)
            
            doc_title = st.text_input("Document Title", uploaded_file.name)
            
            if st.button("Process and Add Document"):
                with st.spinner("Processing document..."):
                    try:
                        file_data = rag_app.process_file(uploaded_file, uploaded_file.name)
                        if rag_app.add_document(doc_title, file_data):
                            st.success("✅ Document successfully processed and added!")
                        else:
                            st.error("❌ Failed to add document.")
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
    
    # Manage Tab
    with doc_tabs[1]:
        st.subheader("Manage Documents")
        docs = get_available_contracts(rag_app)
        
        if docs:
            for doc in docs:
                with st.expander(f"📄 {doc['title']}"):
                    st.write(f"ID: {doc['id']}")
                    st.write(f"Date: {doc['date']}")
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        if st.button("View Details", key=f"view_{doc['id']}"):
                            st.json(rag_app.contract_metadata.get(doc['id'], {}))
                    with col2:
                        if st.button("Generate Report", key=f"report_{doc['id']}"):
                            with st.spinner("Generating report..."):
                                report_bytes = rag_app.export_legal_report(doc['id'])
                                st.download_button(
                                    "Download Report",
                                    report_bytes,
                                    f"report_{doc['id']}.docx"
                                )
                    with col3:
                        if st.button("Delete", key=f"delete_{doc['id']}"):
                            # Implement delete functionality
                            st.warning("Delete functionality to be implemented")
        else:
            st.info("No documents found. Upload some documents to get started!")
    
    # Templates Tab
    with doc_tabs[2]:
        st.subheader("Document Templates")
        st.info("Document template management coming soon!")

def show_search_analysis(rag_app):
    """Search and analysis interface"""
    st.header("🔍 Search & Analysis")
    
    # Create tabs for different search operations
    search_tabs = st.tabs(["Quick Search", "Advanced Search", "Analysis"])
    
    # Quick Search Tab
    with search_tabs[0]:
        query = st.text_input("Enter your search query:")
        if st.button("Search"):
            if query:
                with st.spinner("Searching..."):
                    results = rag_app.search_documents(query)
                    if results:
                        for doc in results:
                            with st.expander(f"📄 {doc['title']}"):
                                st.write(doc['content'])
                                st.write("---")
                                st.write(doc.get('metadata', ''))
                    else:
                        st.warning("No results found")
    
    # Advanced Search Tab
    with search_tabs[1]:
        st.subheader("Advanced Search")
        
        # Search filters
        col1, col2 = st.columns(2)
        with col1:
            doc_type = st.selectbox(
                "Document Type",
                ["All", "Contract", "Invoice", "Legal Brief", "Other"]
            )
        with col2:
            date_range = st.date_input(
                "Date Range",
                value=[],
                key="date_range"
            )
        
        # Additional filters
        filters = {}
        if doc_type != "All":
            filters["document_type"] = doc_type
        if date_range:
            filters["date_range"] = {
                "start": date_range[0].isoformat() if len(date_range) > 0 else None,
                "end": date_range[1].isoformat() if len(date_range) > 1 else None
            }
        
        advanced_query = st.text_input("Advanced Search Query:")
        if st.button("Advanced Search"):
            if advanced_query:
                with st.spinner("Searching..."):
                    results = rag_app.search_documents(advanced_query, filters=filters)
                    if results:
                        for doc in results:
                            with st.expander(f"📄 {doc['title']}"):
                                st.write(doc['content'])
                                st.write("---")
                                st.write(doc.get('metadata', ''))
                    else:
                        st.warning("No results found")
    
    # Analysis Tab
    with search_tabs[2]:
        st.subheader("Document Analysis")
        docs = get_available_contracts(rag_app)
        
        selected_docs = st.multiselect(
            "Select Documents for Analysis",
            docs,
            format_func=lambda x: x['title']
        )
        
        if selected_docs:
            if st.button("Analyze"):
                with st.spinner("Analyzing documents..."):
                    for doc in selected_docs:
                        with st.expander(f"Analysis: {doc['title']}"):
                            try:
                                analysis = rag_app.get_legal_summary(doc['id'])
                                st.json(analysis)
                            except Exception as e:
                                st.error(f"Error analyzing document: {str(e)}")

def show_settings(rag_app):
    """Settings interface"""
    st.header("⚙️ Settings")
    
    # Create tabs for different settings
    settings_tabs = st.tabs(["General", "API Keys", "Advanced"])
    
    # General Settings Tab
    with settings_tabs[0]:
        st.subheader("General Settings")
        
        # Display settings
        st.write("Display Settings")
        st.checkbox("Dark Mode", value=True)
        st.slider("Results per page", min_value=5, max_value=50, value=10)
        
        # Processing settings
        st.write("Processing Settings")
        st.number_input("Maximum file size (MB)", min_value=1, max_value=100, value=10)
        st.selectbox("Default document type", ["Contract", "Invoice", "Legal Brief", "Other"])
    
    # API Keys Tab
    with settings_tabs[1]:
        st.subheader("API Keys")
        
        # Mask the actual keys
        st.text_input("OpenAI API Key", value="*****", type="password")
        st.text_input("Weaviate API Key", value="*****", type="password")
        st.text_input("Weaviate URL", value="*****", type="password")
        
        if st.button("Update Keys"):
            st.warning("Key update functionality to be implemented")
    
    # Advanced Settings Tab
    with settings_tabs[2]:
        st.subheader("Advanced Settings")
        
        # Model settings
        st.write("Model Settings")
        st.selectbox("GPT Model", ["gpt-4", "gpt-3.5-turbo"])
        st.slider("Temperature", min_value=0.0, max_value=1.0, value=0.7)
        
        # Database settings
        st.write("Database Settings")
        if st.button("Backup Database"):
            try:
                with st.spinner("Creating backup..."):
                    if rag_app.backup_database("backup.json"):
                        st.success("Backup created successfully!")
                    else:
                        st.error("Backup failed")
            except Exception as e:
                st.error(f"Error creating backup: {str(e)}")
        
        if st.button("Restore Database"):
            st.warning("Restore functionality to be implemented")
def show_reports_export(rag_app):
    st.header("📊 Reports & Export")
    
    # Create tabs for different report types
    report_tabs = st.tabs([
        "Contract Reports", 
        "Compliance Analysis", 
        "Due Diligence", 
        "Audit Trails",
        "Batch Export"
    ])
    
    # Contract Reports Tab
    with report_tabs[0]:
        st.subheader("📄 Contract Reports")
        
        # Contract selection
        contracts = get_available_contracts(rag_app)  # Implement this helper function
        selected_contract = st.selectbox(
            "Select Contract",
            contracts,
            format_func=lambda x: f"{x['title']} ({x['id']})"
        )
        
        # Report template selection
        template_type = st.selectbox(
            "Report Template",
            ["detailed_analysis", "executive_summary", "compliance_report"],
            format_func=lambda x: x.replace("_", " ").title()
        )
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("Generate Report"):
                with st.spinner("Generating report..."):
                    try:
                        report_bytes = rag_app.export_legal_report(
                            selected_contract['id'],
                            template_type
                        )
                        
                        # Create download button
                        st.download_button(
                            label="Download Report",
                            data=report_bytes,
                            file_name=f"legal_report_{selected_contract['id']}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        st.success("Report generated successfully!")
                    except Exception as e:
                        st.error(f"Error generating report: {str(e)}")
        
        with col2:
            if st.button("Preview Report"):
                with st.spinner("Generating preview..."):
                    try:
                        # Get report data for preview
                        preview_data = {
                            "legal_summary": rag_app.get_formatted_legal_summary(selected_contract['id']),
                            "due_diligence": rag_app.get_due_diligence_report(selected_contract['id'])
                        }
                        
                        # Show preview in expandable sections
                        for section, data in preview_data.items():
                            with st.expander(section.replace("_", " ").title()):
                                st.json(data)
                    except Exception as e:
                        st.error(f"Error generating preview: {str(e)}")
    
    # Compliance Analysis Tab
    with report_tabs[1]:
        st.subheader("⚖️ Compliance Analysis")
        
        # Multiple contract selection for compliance
        selected_contracts = st.multiselect(
            "Select Contracts for Compliance Analysis",
            contracts,
            format_func=lambda x: f"{x['title']} ({x['id']})"
        )
        
        if st.button("Run Compliance Analysis"):
            with st.spinner("Analyzing compliance..."):
                for contract in selected_contracts:
                    try:
                        compliance_report = rag_app.get_compliance_report(contract['id'])
                        with st.expander(f"Compliance Report - {contract['title']}"):
                            st.json(compliance_report)
                            
                            # Add export button for individual compliance report
                            report_bytes = rag_app.export_legal_report(
                                contract['id'],
                                "compliance_report"
                            )
                            st.download_button(
                                label="Download Compliance Report",
                                data=report_bytes,
                                file_name=f"compliance_report_{contract['id']}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    except Exception as e:
                        st.error(f"Error analyzing compliance for {contract['title']}: {str(e)}")
    
    # Due Diligence Tab
    with report_tabs[2]:
        st.subheader("🔍 Due Diligence")
        
        # Contract selection for due diligence
        selected_contract = st.selectbox(
            "Select Contract for Due Diligence",
            contracts,
            format_func=lambda x: f"{x['title']} ({x['id']})",
            key="due_diligence_contract"
        )
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("Generate Due Diligence Report"):
                with st.spinner("Generating due diligence report..."):
                    try:
                        due_diligence = rag_app.get_due_diligence_report(selected_contract['id'])
                        
                        # Create report tabs
                        dd_tabs = st.tabs([
                            "Overview",
                            "Risk Analysis",
                            "Recommendations",
                            "Export"
                        ])
                        
                        with dd_tabs[0]:
                            st.json(due_diligence.get("overview", {}))
                        
                        with dd_tabs[1]:
                            st.json(due_diligence.get("risk_analysis", {}))
                        
                        with dd_tabs[2]:
                            st.json(due_diligence.get("recommendations", {}))
                        
                        with dd_tabs[3]:
                            report_bytes = rag_app.export_legal_report(
                                selected_contract['id'],
                                "detailed_analysis"
                            )
                            st.download_button(
                                label="Download Due Diligence Report",
                                data=report_bytes,
                                file_name=f"due_diligence_{selected_contract['id']}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    except Exception as e:
                        st.error(f"Error generating due diligence report: {str(e)}")
    
    # Audit Trails Tab
    with report_tabs[3]:
        st.subheader("📋 Audit Trails")
        
        # Date range selection
        col1, col2 = st.columns(2)
        with col1:
            start_date = st.date_input("Start Date")
        with col2:
            end_date = st.date_input("End Date")
        
        # Contract selection for audit
        selected_contract = st.selectbox(
            "Select Contract for Audit Trail",
            contracts,
            format_func=lambda x: f"{x['title']} ({x['id']})",
            key="audit_contract"
        )
        
        if st.button("Generate Audit Trail"):
            with st.spinner("Generating audit trail..."):
                try:
                    audit_trail = rag_app.get_contract_audit_trail(selected_contract['id'])
                    
                    # Display audit trail
                    st.write("### Activity Timeline")
                    for activity in audit_trail.get("activities", []):
                        st.write(f"**{activity['timestamp']}**: {activity['description']}")
                    
                    # Export option
                    report_bytes = rag_app.export_legal_report(
                        selected_contract['id'],
                        "detailed_analysis"
                    )
                    st.download_button(
                        label="Download Audit Report",
                        data=report_bytes,
                        file_name=f"audit_trail_{selected_contract['id']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Error generating audit trail: {str(e)}")
    
    # Batch Export Tab
    with report_tabs[4]:
        st.subheader("📦 Batch Export")
        
        # Contract selection for batch export
        selected_contracts = st.multiselect(
            "Select Contracts for Batch Export",
            contracts,
            format_func=lambda x: f"{x['title']} ({x['id']})",
            key="batch_export"
        )
        
        # Report type selection
        report_types = st.multiselect(
            "Select Report Types",
            ["Legal Summary", "Due Diligence", "Compliance", "Audit Trail"],
            default=["Legal Summary"]
        )
        
        if st.button("Generate Batch Export"):
            with st.spinner("Generating batch export..."):
                try:
                    # Create a ZIP file containing all reports
                    zip_buffer = generate_batch_export(rag_app, selected_contracts, report_types)
                    
                    st.download_button(
                        label="Download Batch Export",
                        data=zip_buffer.getvalue(),
                        file_name="legal_reports_batch.zip",
                        mime="application/zip"
                    )
                    
                    st.success("Batch export generated successfully!")
                except Exception as e:
                    st.error(f"Error generating batch export: {str(e)}")

def generate_batch_export(rag_app, contracts, report_types):
    """Helper function to generate batch export"""
    import io
    import zipfile
    
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for contract in contracts:
            for report_type in report_types:
                try:
                    if report_type == "Legal Summary":
                        report_bytes = rag_app.export_legal_report(
                            contract['id'],
                            "executive_summary"
                        )
                    elif report_type == "Due Diligence":
                        report_bytes = rag_app.export_legal_report(
                            contract['id'],
                            "detailed_analysis"
                        )
                    elif report_type == "Compliance":
                        report_bytes = rag_app.export_legal_report(
                            contract['id'],
                            "compliance_report"
                        )
                    else:  # Audit Trail
                        audit_trail = rag_app.get_contract_audit_trail(contract['id'])
                        report_bytes = rag_app.export_legal_report(
                            contract['id'],
                            "detailed_analysis"
                        )
                    
                    filename = f"{contract['id']}_{report_type.lower()}.docx"
                    zip_file.writestr(filename, report_bytes)
                    
                except Exception as e:
                    print(f"Error processing {report_type} for {contract['id']}: {str(e)}")
    
    return zip_buffer

def get_available_contracts(rag_app):
    """Helper function to get available contracts"""
    # Implement this based on your data structure
    contracts = []
    for contract_id, metadata in rag_app.contract_metadata.items():
        contracts.append({
            'id': contract_id,
            'title': metadata.get('filename', 'Unnamed Contract'),
            'date': metadata.get('processing_date', 'Unknown')
        })
    return contracts

if __name__ == "__main__":
    main()